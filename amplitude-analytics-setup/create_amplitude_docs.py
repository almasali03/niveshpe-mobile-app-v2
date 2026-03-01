#!/usr/bin/env python3
"""
Generate Amplitude Analytics documentation for NiveshPe React Native App
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_amplitude_setup_doc():
    # Create a new Document
    doc = Document()
    
    # Set document title
    title = doc.add_heading('Amplitude Analytics Setup Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_heading('NiveshPe React Native App (iOS & Android)', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Installation & Setup',
        '2. SDK Configuration',
        '3. Authentication Tracking',
        '4. Screen Navigation Tracking',
        '5. Investment & Revenue Tracking',
        '6. Push Notifications & Deep Links',
        '7. Error & Performance Monitoring',
        '8. Testing & Debugging',
        '9. Privacy & Compliance',
        '10. Implementation Checklist'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 1: Installation
    doc.add_heading('1. Installation & Setup', level=1)
    
    doc.add_heading('Install Dependencies', level=2)
    doc.add_paragraph('Install the Amplitude React Native SDK and required dependencies:')
    
    code_block = doc.add_paragraph()
    code_block.add_run('# Using npm\n').font.name = 'Courier New'
    code_block.add_run('npm install @amplitude/react-native@^2.17.0\n').font.name = 'Courier New'
    code_block.add_run('npm install react-native-device-info\n').font.name = 'Courier New'
    code_block.add_run('npm install @react-native-async-storage/async-storage\n\n').font.name = 'Courier New'
    code_block.add_run('# For iOS\n').font.name = 'Courier New'
    code_block.add_run('cd ios && pod install').font.name = 'Courier New'
    
    doc.add_heading('iOS Additional Setup', level=2)
    doc.add_paragraph('Add to your iOS Info.plist for tracking transparency:')
    
    ios_code = doc.add_paragraph()
    ios_code.add_run('<key>NSUserTrackingUsageDescription</key>\n').font.name = 'Courier New'
    ios_code.add_run('<string>This app uses tracking to improve your investment experience.</string>').font.name = 'Courier New'
    
    doc.add_page_break()
    
    # Section 2: SDK Configuration
    doc.add_heading('2. SDK Configuration', level=1)
    
    doc.add_heading('Create Amplitude Service', level=2)
    doc.add_paragraph('Create a new file: src/services/amplitude.js')
    
    doc.add_paragraph('Key Configuration Points:')
    config_points = [
        'Use separate API keys for development and production',
        'Initialize on app launch',
        'Set user properties from device info',
        'Track app lifecycle events',
        'Handle errors gracefully'
    ]
    for point in config_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('Environment Variables', level=2)
    doc.add_paragraph('Create .env files for different environments:')
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Environment'
    hdr_cells[1].text = 'Configuration'
    
    # Data rows
    table.rows[1].cells[0].text = '.env.development'
    table.rows[1].cells[1].text = 'AMPLITUDE_API_KEY=dev_key_here'
    
    table.rows[2].cells[0].text = '.env.production'
    table.rows[2].cells[1].text = 'AMPLITUDE_API_KEY=prod_key_here'
    
    doc.add_page_break()
    
    # Section 3: Authentication Tracking
    doc.add_heading('3. Authentication Tracking', level=1)
    
    doc.add_heading('Key Authentication Events', level=2)
    
    auth_table = doc.add_table(rows=7, cols=3)
    auth_table.style = 'Light Grid Accent 1'
    
    # Header
    hdr = auth_table.rows[0].cells
    hdr[0].text = 'Event Name'
    hdr[1].text = 'Trigger Point'
    hdr[2].text = 'Key Properties'
    
    # Events
    events = [
        ('phone_number_entered', 'User enters phone', 'page_source, entry_point'),
        ('otp_requested', 'OTP sent', 'phone_masked, request_time'),
        ('otp_verified', 'OTP verified', 'attempts, verification_time'),
        ('pan_submitted', 'PAN entered', 'pan_masked, user_type'),
        ('pan_verified', 'PAN verified', 'existing_investor, name_match'),
        ('kyc_completed', 'KYC done', 'completion_time, user_type')
    ]
    
    for i, (event, trigger, props) in enumerate(events, 1):
        row = auth_table.rows[i].cells
        row[0].text = event
        row[1].text = trigger
        row[2].text = props
    
    doc.add_heading('Implementation Example', level=2)
    doc.add_paragraph('After OTP verification, set user ID and properties:')
    
    code_example = doc.add_paragraph()
    code_example.add_run('// Set user ID after OTP verification\n').font.name = 'Courier New'
    code_example.add_run('await AmplitudeService.setUserId(phoneNumber);\n').font.name = 'Courier New'
    code_example.add_run('await AmplitudeService.setUserProperties({\n').font.name = 'Courier New'
    code_example.add_run('  user_type: "new",\n').font.name = 'Courier New'
    code_example.add_run('  kyc_status: "pending"\n').font.name = 'Courier New'
    code_example.add_run('});').font.name = 'Courier New'
    
    doc.add_page_break()
    
    # Section 4: Screen Tracking
    doc.add_heading('4. Screen Navigation Tracking', level=1)
    
    doc.add_heading('Automatic Screen Tracking', level=2)
    doc.add_paragraph('Use React Navigation integration for automatic screen tracking:')
    
    screens = [
        'Dashboard', 'Mutual Funds', 'Fund Details', 'Add Funds',
        'My Portfolio', 'Goals', 'Profile', 'KYC Screens',
        'Withdrawal', 'Transactions', 'Calculators', 'Support'
    ]
    
    doc.add_paragraph('Key screens to track:')
    for screen in screens:
        doc.add_paragraph(f'• {screen}', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 5: Investment Tracking
    doc.add_heading('5. Investment & Revenue Tracking', level=1)
    
    doc.add_heading('Critical Investment Events', level=2)
    
    inv_table = doc.add_table(rows=6, cols=3)
    inv_table.style = 'Light Grid Accent 1'
    
    hdr = inv_table.rows[0].cells
    hdr[0].text = 'Event'
    hdr[1].text = 'When to Track'
    hdr[2].text = 'Revenue Impact'
    
    inv_events = [
        ('fund_selected', 'User selects fund', 'No'),
        ('investment_amount_entered', 'Amount entered', 'No'),
        ('payment_method_selected', 'Payment chosen', 'No'),
        ('investment_completed', 'Payment successful', 'Yes - Track as revenue'),
        ('sip_activated', 'SIP started', 'Yes - Track monthly')
    ]
    
    for i, (event, when, revenue) in enumerate(inv_events, 1):
        row = inv_table.rows[i].cells
        row[0].text = event
        row[1].text = when
        row[2].text = revenue
    
    doc.add_heading('Revenue Tracking', level=2)
    doc.add_paragraph('Track all successful investments as revenue:')
    doc.add_paragraph('• SIP investments: Track monthly amount')
    doc.add_paragraph('• Lumpsum investments: Track full amount')
    doc.add_paragraph('• Goal-based investments: Track with goal context')
    
    doc.add_page_break()
    
    # Section 6: Push & Deep Links
    doc.add_heading('6. Push Notifications & Deep Links', level=1)
    
    doc.add_heading('Push Notification Events', level=2)
    notif_events = [
        'push_notification_received',
        'push_notification_opened',
        'push_notification_dismissed',
        'push_permission_granted',
        'push_permission_denied'
    ]
    
    for event in notif_events:
        doc.add_paragraph(f'• {event}', style='List Bullet')
    
    doc.add_heading('Deep Link Tracking', level=2)
    doc.add_paragraph('Track all deep link opens with:')
    doc.add_paragraph('• Source URL')
    doc.add_paragraph('• Campaign parameters')
    doc.add_paragraph('• Navigation destination')
    
    doc.add_page_break()
    
    # Section 7: Error Tracking
    doc.add_heading('7. Error & Performance Monitoring', level=1)
    
    doc.add_heading('Error Types to Track', level=2)
    error_table = doc.add_table(rows=5, cols=2)
    error_table.style = 'Light Grid Accent 1'
    
    hdr = error_table.rows[0].cells
    hdr[0].text = 'Error Type'
    hdr[1].text = 'Details to Capture'
    
    errors = [
        ('JavaScript Errors', 'Stack trace, error message, screen'),
        ('API Failures', 'Endpoint, status code, error response'),
        ('Payment Failures', 'Payment method, amount, error code'),
        ('Network Errors', 'Connection type, retry count, endpoint')
    ]
    
    for i, (error_type, details) in enumerate(errors, 1):
        row = error_table.rows[i].cells
        row[0].text = error_type
        row[1].text = details
    
    doc.add_heading('Performance Metrics', level=2)
    doc.add_paragraph('Track these performance indicators:')
    perf_metrics = [
        'Screen load time',
        'API response time',
        'Image load time',
        'App launch time',
        'Time to interactive'
    ]
    for metric in perf_metrics:
        doc.add_paragraph(f'• {metric}', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 8: Testing
    doc.add_heading('8. Testing & Debugging', level=1)
    
    doc.add_heading('Testing Checklist', level=2)
    test_items = [
        'Verify Amplitude initializes on app launch',
        'Check user ID is set after authentication',
        'Confirm events appear in Amplitude dashboard',
        'Test revenue tracking for investments',
        'Verify user properties are updated',
        'Check error tracking works',
        'Test deep link tracking',
        'Verify push notification events'
    ]
    
    for item in test_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐ ' + item)
    
    doc.add_heading('Debug Mode', level=2)
    doc.add_paragraph('Enable verbose logging in development:')
    debug_code = doc.add_paragraph()
    debug_code.add_run('if (__DEV__) {\n').font.name = 'Courier New'
    debug_code.add_run('  Amplitude.getInstance().setLogLevel(LogLevel.VERBOSE);\n').font.name = 'Courier New'
    debug_code.add_run('}').font.name = 'Courier New'
    
    doc.add_page_break()
    
    # Section 9: Privacy
    doc.add_heading('9. Privacy & Compliance', level=1)
    
    doc.add_heading('Data Privacy Guidelines', level=2)
    privacy_guidelines = [
        'Never track full PAN numbers',
        'Mask phone numbers in events',
        'Don\'t log sensitive bank details',
        'Implement user opt-out option',
        'Support data deletion requests',
        'Comply with local privacy laws'
    ]
    
    for guideline in privacy_guidelines:
        doc.add_paragraph(f'• {guideline}', style='List Bullet')
    
    doc.add_heading('User Consent', level=2)
    doc.add_paragraph('Implement consent flow:')
    doc.add_paragraph('1. Show tracking consent on first launch')
    doc.add_paragraph('2. Allow users to opt-out in settings')
    doc.add_paragraph('3. Respect user preferences')
    
    doc.add_page_break()
    
    # Section 10: Implementation Checklist
    doc.add_heading('10. Implementation Checklist', level=1)
    
    doc.add_heading('Setup Phase', level=2)
    setup_tasks = [
        'Create Amplitude account',
        'Get API keys (dev & prod)',
        'Install React Native SDK',
        'Configure environment variables',
        'Set up Amplitude service class'
    ]
    
    for task in setup_tasks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐ ' + task)
    
    doc.add_heading('Implementation Phase', level=2)
    impl_tasks = [
        'Add screen tracking',
        'Implement authentication events',
        'Add investment tracking',
        'Set up revenue tracking',
        'Implement error tracking',
        'Add push notification events',
        'Set up deep link tracking'
    ]
    
    for task in impl_tasks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐ ' + task)
    
    doc.add_heading('Testing Phase', level=2)
    test_tasks = [
        'Test in development environment',
        'Verify events in dashboard',
        'Test user properties',
        'Validate revenue tracking',
        'Test error scenarios',
        'Performance testing'
    ]
    
    for task in test_tasks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐ ' + task)
    
    doc.add_heading('Launch Phase', level=2)
    launch_tasks = [
        'Switch to production API key',
        'Deploy to TestFlight/Beta',
        'Monitor initial data',
        'Create dashboards',
        'Set up alerts',
        'Train team on analytics'
    ]
    
    for task in launch_tasks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐ ' + task)
    
    # Add appendix
    doc.add_page_break()
    doc.add_heading('Appendix: Quick Reference', level=1)
    
    doc.add_heading('Common Event Names', level=2)
    common_events = [
        'app_opened', 'session_started', 'screen_viewed',
        'phone_number_entered', 'otp_verified', 'kyc_completed',
        'fund_viewed', 'investment_completed', 'goal_created',
        'referral_clicked', 'calculator_used', 'error_occurred'
    ]
    
    for event in common_events:
        doc.add_paragraph(f'• {event}', style='List Bullet')
    
    doc.add_heading('Key User Properties', level=2)
    user_props = [
        'user_id (phone number)',
        'user_type (new/existing)',
        'kyc_status',
        'total_invested',
        'portfolio_value',
        'goals_count',
        'risk_profile',
        'app_version'
    ]
    
    for prop in user_props:
        doc.add_paragraph(f'• {prop}', style='List Bullet')
    
    # Save the document
    doc.save('/Users/almas/niveshpe-mobile-app-v1/amplitude-analytics-setup/Amplitude_Setup_Guide_ReactNative.docx')
    print("✅ Word document created: Amplitude_Setup_Guide_ReactNative.docx")

def create_implementation_code_files():
    """Create implementation code files for React Native"""
    
    # 1. Amplitude Service
    amplitude_service = '''import * as Amplitude from '@amplitude/react-native';
import DeviceInfo from 'react-native-device-info';
import AsyncStorage from '@react-native-async-storage/async-storage';

const AMPLITUDE_API_KEY = __DEV__ 
  ? 'YOUR_DEV_API_KEY_HERE'
  : 'YOUR_PROD_API_KEY_HERE';

class AmplitudeService {
  initialized = false;

  async initialize() {
    if (this.initialized) return;
    
    try {
      await Amplitude.init(AMPLITUDE_API_KEY);
      await this.setInitialUserProperties();
      await this.trackAppOpen();
      this.initialized = true;
      console.log('Amplitude initialized successfully');
    } catch (error) {
      console.error('Error initializing Amplitude:', error);
    }
  }

  async setInitialUserProperties() {
    const deviceInfo = {
      device_id: DeviceInfo.getUniqueId(),
      device_brand: DeviceInfo.getBrand(),
      device_model: DeviceInfo.getModel(),
      system_name: DeviceInfo.getSystemName(),
      system_version: DeviceInfo.getSystemVersion(),
      app_version: DeviceInfo.getVersion(),
      app_build: DeviceInfo.getBuildNumber(),
      is_tablet: DeviceInfo.isTablet(),
    };
    
    await Amplitude.setUserProperties(deviceInfo);
  }

  async trackAppOpen() {
    const lastOpenTime = await AsyncStorage.getItem('last_app_open');
    const timeSinceLastOpen = lastOpenTime 
      ? Date.now() - parseInt(lastOpenTime) 
      : 0;
    
    await this.trackEvent('app_opened', {
      time_since_last_open: timeSinceLastOpen,
      launch_source: 'direct',
    });
    
    await AsyncStorage.setItem('last_app_open', Date.now().toString());
  }

  async trackEvent(eventName, properties = {}) {
    try {
      await Amplitude.logEvent(eventName, {
        ...properties,
        timestamp: new Date().toISOString(),
        platform: DeviceInfo.getSystemName(),
      });
      
      if (__DEV__) {
        console.log(`📊 Event: ${eventName}`, properties);
      }
    } catch (error) {
      console.error(`Error tracking event ${eventName}:`, error);
    }
  }

  async setUserId(userId) {
    await Amplitude.setUserId(userId);
  }

  async setUserProperties(properties) {
    await Amplitude.setUserProperties(properties);
  }

  async trackRevenue(amount, productId, properties = {}) {
    const revenue = new Amplitude.Revenue()
      .setPrice(amount)
      .setQuantity(1)
      .setProductId(productId)
      .setRevenueType('investment');
    
    await Amplitude.logRevenue(revenue);
  }
}

export default new AmplitudeService();
'''
    
    with open('/Users/almas/niveshpe-mobile-app-v1/amplitude-analytics-setup/AmplitudeService.js', 'w') as f:
        f.write(amplitude_service)
    print("✅ Created: AmplitudeService.js")
    
    # 2. Auth Tracking
    auth_tracking = '''import AmplitudeService from '../services/AmplitudeService';
import AsyncStorage from '@react-native-async-storage/async-storage';

export const AuthTracking = {
  async trackPhoneEntry(phoneNumber) {
    await AmplitudeService.trackEvent('phone_number_entered', {
      page_source: 'auth',
      entry_point: 'app_launch',
    });
  },

  async trackOTPRequest(phoneNumber) {
    await AmplitudeService.trackEvent('otp_requested', {
      phone_number_masked: phoneNumber.substring(0, 6) + '****',
      request_time: Date.now(),
    });
  },

  async trackOTPVerified(phoneNumber, attempts = 1) {
    await AmplitudeService.trackEvent('otp_verified', {
      phone_number_masked: phoneNumber.substring(0, 6) + '****',
      attempts,
      verification_time: Date.now(),
    });
    
    await AmplitudeService.setUserId(phoneNumber);
    await AsyncStorage.setItem('user_phone', phoneNumber);
  },

  async trackPANSubmitted(panMasked) {
    await AmplitudeService.trackEvent('pan_submitted', {
      pan_masked: panMasked,
    });
  },

  async trackPANVerified(panMasked, isExistingInvestor) {
    await AmplitudeService.trackEvent('pan_verified', {
      pan_masked: panMasked,
      existing_investor: isExistingInvestor,
    });
    
    await AmplitudeService.setUserProperties({
      pan_verified: true,
      user_type: isExistingInvestor ? 'existing' : 'new',
    });
  },

  async trackKYCCompleted() {
    await AmplitudeService.trackEvent('kyc_completed', {
      completion_time: Date.now(),
    });
    
    await AmplitudeService.setUserProperties({
      kyc_status: 'completed',
      kyc_completion_date: new Date().toISOString(),
    });
  },
};
'''
    
    with open('/Users/almas/niveshpe-mobile-app-v1/amplitude-analytics-setup/AuthTracking.js', 'w') as f:
        f.write(auth_tracking)
    print("✅ Created: AuthTracking.js")
    
    # 3. Investment Tracking
    investment_tracking = '''import AmplitudeService from '../services/AmplitudeService';
import AsyncStorage from '@react-native-async-storage/async-storage';

export const InvestmentTracking = {
  async trackFundViewed(fundName, fundCategory, returns) {
    await AmplitudeService.trackEvent('fund_viewed', {
      fund_name: fundName,
      fund_category: fundCategory,
      fund_returns: returns,
      view_source: 'browse',
    });
  },

  async trackFundSelected(fundName, fundCategory, expectedReturns) {
    await AmplitudeService.trackEvent('fund_selected', {
      fund_name: fundName,
      fund_category: fundCategory,
      expected_returns: expectedReturns,
    });
  },

  async trackInvestmentAmount(amount, investmentType, fundName) {
    await AmplitudeService.trackEvent('investment_amount_entered', {
      amount,
      investment_type: investmentType,
      fund_name: fundName,
    });
  },

  async trackPaymentMethod(method, fundName, amount) {
    await AmplitudeService.trackEvent('payment_method_selected', {
      payment_method: method,
      fund_name: fundName,
      amount,
    });
  },

  async trackInvestmentCompleted(fundName, amount, type, paymentMethod) {
    await AmplitudeService.trackEvent('investment_completed', {
      fund_name: fundName,
      amount,
      investment_type: type,
      payment_method: paymentMethod,
    });
    
    // Track revenue
    await AmplitudeService.trackRevenue(amount, type, {
      fund_name: fundName,
      payment_method: paymentMethod,
    });
    
    // Update user properties
    const totalInvested = await AsyncStorage.getItem('total_invested') || '0';
    const newTotal = parseFloat(totalInvested) + amount;
    await AsyncStorage.setItem('total_invested', newTotal.toString());
    
    await AmplitudeService.setUserProperties({
      total_invested: newTotal,
      last_investment_date: new Date().toISOString(),
    });
  },

  async trackSIPActivated(fundName, monthlyAmount, duration) {
    await AmplitudeService.trackEvent('sip_activated', {
      fund_name: fundName,
      monthly_amount: monthlyAmount,
      duration_months: duration,
      total_commitment: monthlyAmount * duration,
    });
    
    // Track as revenue
    await AmplitudeService.trackRevenue(monthlyAmount, 'sip', {
      fund_name: fundName,
      recurring: true,
    });
  },
};
'''
    
    with open('/Users/almas/niveshpe-mobile-app-v1/amplitude-analytics-setup/InvestmentTracking.js', 'w') as f:
        f.write(investment_tracking)
    print("✅ Created: InvestmentTracking.js")
    
    # 4. Screen Tracking Hook
    screen_tracking = '''import { useEffect } from 'react';
import { useNavigation, useRoute } from '@react-navigation/native';
import AmplitudeService from '../services/AmplitudeService';

export function useScreenTracking(screenName, properties = {}) {
  const navigation = useNavigation();
  const route = useRoute();
  
  useEffect(() => {
    const unsubscribe = navigation.addListener('focus', () => {
      AmplitudeService.trackEvent('screen_viewed', {
        screen_name: screenName,
        previous_screen: navigation.getState()?.routes[navigation.getState().index - 1]?.name,
        params: route.params,
        ...properties,
      });
    });
    
    return unsubscribe;
  }, [navigation, screenName]);
}

// Usage example:
// function DashboardScreen() {
//   useScreenTracking('Dashboard');
//   return <View>...</View>;
// }
'''
    
    with open('/Users/almas/niveshpe-mobile-app-v1/amplitude-analytics-setup/useScreenTracking.js', 'w') as f:
        f.write(screen_tracking)
    print("✅ Created: useScreenTracking.js")

if __name__ == '__main__':
    try:
        # Install python-docx if not already installed
        import subprocess
        import sys
        
        try:
            from docx import Document
        except ImportError:
            print("Installing python-docx...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            from docx import Document
        
        # Create the Word document
        create_amplitude_setup_doc()
        
        # Create implementation files
        create_implementation_code_files()
        
        print("\n✅ All files created successfully in amplitude-analytics-setup folder!")
        print("\n📁 Folder contents:")
        print("  - Amplitude_Setup_Guide_ReactNative.docx (Main documentation)")
        print("  - amplitude_tracking_plan.csv (Event tracking spreadsheet)")
        print("  - AmplitudeService.js (Core service)")
        print("  - AuthTracking.js (Authentication events)")
        print("  - InvestmentTracking.js (Investment events)")
        print("  - useScreenTracking.js (Screen tracking hook)")
        
    except Exception as e:
        print(f"Error: {e}")