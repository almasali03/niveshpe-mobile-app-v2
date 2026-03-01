#!/usr/bin/env python3
"""
iOS App Store ACTUAL Rejection Risks Document Generator
Based on real 2024-2025 App Store rejection data (24.8% rejection rate of 7.77M submissions)
Focus: Only guidelines that ACTUALLY cause rejection, not theoretical best practices
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Color definitions
COLORS = {
    'critical': RGBColor(239, 68, 68),    # Red
    'high': RGBColor(245, 158, 11),       # Orange
    'medium': RGBColor(251, 191, 36),     # Yellow
    'pass': RGBColor(16, 185, 129),       # Green
    'blue': RGBColor(37, 99, 235),        # Blue
}

def create_document():
    """Create and configure the Word document"""
    doc = Document()
    doc.core_properties.title = "iOS App Store ACTUAL Rejection Risks - NiveshPe App"
    doc.core_properties.author = "NiveshPe Development Team"
    return doc

def add_title_page(doc):
    """Add formatted title page"""
    title = doc.add_paragraph()
    title_run = title.add_run("iOS App Store\nACTUAL Rejection Risks")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = COLORS['blue']
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("NiveshPe Mutual Fund Investment App")
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_run = date_para.add_run("December 2024 | Based on 2024-2025 Rejection Data")
    date_run.font.size = Pt(14)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info_run = info.add_run(
        "Focus: Only design issues that ACTUALLY cause App Store rejection\n"
        "Data Source: 24.8% rejection rate of 7.77M submissions in 2024\n"
        "App Analyzed: NiveshPe (39 HTML pages, mutual fund investment platform)\n"
        "Priority: Critical bugs, compliance issues, and quality violations"
    )
    info_run.font.size = Pt(11)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

def add_executive_summary(doc):
    """Add executive summary"""
    doc.add_heading('EXECUTIVE SUMMARY', 0)

    doc.add_heading('What This Document Covers', 1)

    intro = doc.add_paragraph(
        "This document identifies design and compliance issues that have a REAL risk of "
        "causing App Store rejection based on 2024-2025 data. We do NOT cover theoretical "
        "best practices or accessibility features that are optional."
    )
    intro_run = intro.runs[0]
    intro_run.font.size = Pt(11)

    doc.add_paragraph()

    # Key findings
    doc.add_heading('Key Findings for NiveshPe App', 2)

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = "Issue"
    headers[1].text = "Rejection Risk"
    headers[2].text = "Status"

    # Make headers bold
    for cell in headers:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Row 1: Crashes/Stability
    row1 = table.rows[1].cells
    row1[0].text = "App Crashes or Technical Instability"
    row1[1].text = "CRITICAL - Automatic Rejection"
    row1[1].paragraphs[0].runs[0].font.color.rgb = COLORS['critical']
    row1[2].text = "⚠️ Needs Testing"

    # Row 2: Incomplete Build
    row2 = table.rows[2].cells
    row2[0].text = "Incomplete Build / Placeholder Content"
    row2[1].text = "CRITICAL - 40% of Rejections"
    row2[1].paragraphs[0].runs[0].font.color.rgb = COLORS['critical']
    row2[2].text = "✓ Appears Complete"

    # Row 3: Financial Compliance
    row3 = table.rows[3].cells
    row3[0].text = "Financial Institution Verification (3.2.1)"
    row3[1].text = "CRITICAL - Fintech Apps"
    row3[1].paragraphs[0].runs[0].font.color.rgb = COLORS['critical']
    row3[2].text = "❌ SEBI Docs Needed"

    # Row 4: Privacy Manifest
    row4 = table.rows[4].cells
    row4[0].text = "Privacy Manifest Missing (Required May 2024)"
    row4[1].text = "CRITICAL - Automatic Rejection"
    row4[1].paragraphs[0].runs[0].font.color.rgb = COLORS['critical']
    row4[2].text = "❌ Not Implemented"

    # Row 5: Native iOS Design Violations (NEW)
    row5 = table.rows[5].cells
    row5[0].text = "Non-Native iOS Design (Material Design/Android)"
    row5[1].text = "HIGH - Looks Like Web/Android App"
    row5[1].paragraphs[0].runs[0].font.color.rgb = COLORS['high']
    row5[2].text = "❌ Uses Material Design"

    # Row 6: Poor UI Quality
    row6 = table.rows[6].cells
    row6[0].text = "Poor UI Quality / Confusing Navigation"
    row6[1].text = "HIGH - Subjective Rejection"
    row6[1].paragraphs[0].runs[0].font.color.rgb = COLORS['high']
    row6[2].text = "⚠️ Needs Review"

    # Row 7: Broken Links
    row7 = table.rows[7].cells
    row7[0].text = "Broken Links / Non-functional Elements"
    row7[1].text = "MEDIUM-HIGH"
    row7[1].paragraphs[0].runs[0].font.color.rgb = COLORS['medium']
    row7[2].text = "⚠️ Verify All Links"

    doc.add_paragraph()

    # Timeline
    doc.add_heading('Immediate Action Required', 2)

    actions = [
        "1. Device testing to ensure no crashes (mandatory before submission)",
        "2. SEBI registration documentation for financial app verification",
        "3. Implement Privacy Manifest (PrivacyInfo.xcprivacy) - Required since May 2024",
        "4. Replace Material Design with native iOS design patterns (SF Symbols, native navigation)",
        "5. Remove custom status bar, use native iOS status bar",
        "6. Remove any placeholder content or broken links",
        "7. Test on multiple devices and iOS versions (17+)",
        "8. Prepare demo account credentials for App Review team"
    ]

    for action in actions:
        p = doc.add_paragraph(action, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_critical_section_crashes(doc):
    """Section 1: Crashes and Technical Stability"""
    doc.add_heading('SECTION 1: CRASHES & TECHNICAL STABILITY', 0)

    doc.add_heading('Rejection Risk: CRITICAL (Automatic Rejection)', 1)
    risk = doc.add_paragraph()
    risk_run = risk.add_run("⛔ GUARANTEED REJECTION IF APP CRASHES DURING REVIEW")
    risk_run.font.bold = True
    risk_run.font.size = Pt(12)
    risk_run.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    doc.add_heading('Apple Guideline Reference', 2)
    doc.add_paragraph(
        "• Guideline 2.1: App Completeness - 'Over 40% of rejections relate to this guideline'\n"
        "• Apps must be stable, bug-free, and thoroughly tested on actual devices\n"
        "• Performance issues (slow loading, unresponsive UI) also cause rejection"
    )

    doc.add_heading('What Apple Tests For', 2)

    tests = [
        "App launches without crashing",
        "All core features work as described",
        "No freezing or unresponsive screens",
        "Smooth navigation between pages",
        "Fast loading times (no unusually slow performance)",
        "No broken UI elements or overlapping content"
    ]

    for test in tests:
        doc.add_paragraph(f"✓ {test}", style='List Bullet')

    doc.add_heading('Current Status: NiveshPe App', 2)

    status_critical = doc.add_paragraph()
    status_critical.add_run("⚠️ STATUS: UNKNOWN - REQUIRES DEVICE TESTING").font.bold = True
    status_critical.add_run("\n\nSince this is a web-based prototype (HTML/CSS/JS), you MUST test on actual iOS devices using:")

    testing_methods = [
        "WKWebView wrapper (for web apps)",
        "Capacitor or Cordova (if converting to native)",
        "iOS Safari on multiple devices (iPhone SE, iPhone 14, iPhone 15 Pro)",
        "Different iOS versions (iOS 17.0, 17.5, 18.0+)"
    ]

    for method in testing_methods:
        doc.add_paragraph(f"• {method}", style='List Bullet')

    doc.add_heading('Common Crash Causes in Web-to-iOS Apps', 2)

    crashes = [
        ("Memory Issues", "Large images, infinite scrolling without cleanup"),
        ("JavaScript Errors", "Uncaught exceptions, undefined variables"),
        ("Missing Resources", "404 errors for CSS/JS files, broken image links"),
        ("iOS-Specific Bugs", "Input focus issues, scroll behavior, touch events"),
        ("Network Failures", "API calls without error handling, timeout issues")
    ]

    crash_table = doc.add_table(rows=len(crashes)+1, cols=2)
    crash_table.style = 'Light List Accent 1'

    crash_headers = crash_table.rows[0].cells
    crash_headers[0].text = "Crash Type"
    crash_headers[1].text = "Common Causes in NiveshPe App"
    for cell in crash_headers:
        cell.paragraphs[0].runs[0].font.bold = True

    for i, (crash_type, cause) in enumerate(crashes, 1):
        row = crash_table.rows[i].cells
        row[0].text = crash_type
        row[1].text = cause

    doc.add_paragraph()

    doc.add_heading('Action Items (MANDATORY Before Submission)', 2)

    actions = [
        "□ Test app on physical iPhone devices (not just simulator)",
        "□ Test all 39 pages for crashes and freezes",
        "□ Use iOS Console to monitor JavaScript errors",
        "□ Test with poor network conditions (airplane mode, 3G)",
        "□ Memory profiling (Safari Web Inspector)",
        "□ Test form submissions, OTP verification, fund purchases end-to-end",
        "□ Fix ANY crashes or errors before submission"
    ]

    for action in actions:
        p = doc.add_paragraph(action)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_critical_section_financial(doc):
    """Section 2: Financial App Compliance"""
    doc.add_heading('SECTION 2: FINANCIAL APP COMPLIANCE', 0)

    doc.add_heading('Rejection Risk: CRITICAL (Financial Apps Have Additional Requirements)', 1)
    risk = doc.add_paragraph()
    risk_run = risk.add_run("⛔ FINTECH APPS FACE STRICTER SCRUTINY - 1.7M APPS REJECTED IN 2023")
    risk_run.font.bold = True
    risk_run.font.size = Pt(12)
    risk_run.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    doc.add_heading('Apple Guideline 3.2.1(viii): Financial Services', 2)

    doc.add_paragraph(
        "Apps facilitating trading in financial instruments must be from the financial "
        "institution performing such services or must use a public API offered by the institution."
    )

    requirement = doc.add_paragraph()
    requirement.add_run("REQUIREMENT: You must provide proof that NiveshPe is either:\n").font.bold = True

    options = [
        "1. A SEBI-registered mutual fund distributor (ARN holder), OR",
        "2. A SEBI-registered Investment Advisor (RIA), OR",
        "3. Using official APIs from SEBI-registered entities"
    ]

    for option in options:
        p = doc.add_paragraph(option)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading('Documentation Needed for App Review', 2)

    docs_needed = [
        "SEBI Registration Certificate (ARN or RIA certificate)",
        "Company incorporation documents",
        "Terms of service clearly stating regulatory compliance",
        "Privacy policy with India-specific financial data handling",
        "Risk disclosure statements (required by SEBI)",
        "Customer support contact information (phone, email, physical address)"
    ]

    doc.add_paragraph("Submit these documents in App Store Connect > App Review Information:")
    for doc_item in docs_needed:
        doc.add_paragraph(f"✓ {doc_item}", style='List Bullet')

    doc.add_heading('Current Status: NiveshPe App', 2)

    status = doc.add_paragraph()
    status.add_run("❌ STATUS: MISSING SEBI REGISTRATION DISCLOSURE\n\n").font.color.rgb = COLORS['critical']
    status.add_run(
        "The app does not display SEBI registration number or regulatory disclaimers "
        "anywhere in the interface. This is required by both Apple AND SEBI regulations."
    )

    doc.add_heading('Required Fixes', 2)

    doc.add_paragraph("1. Add SEBI Registration in App Footer (All Pages):")

    code1 = doc.add_paragraph(
        "<!-- Add to footer in index.html and all pages -->\n"
        '<div class="sebi-registration">\n'
        '  <p class="reg-text">\n'
        '    SEBI Registered Investment Advisor: INA000XXXXXX<br>\n'
        '    Compliance Officer: [Name] | Email: compliance@niveshpe.com<br>\n'
        '    Grievance: support@niveshpe.com | 1800-XXX-XXXX\n'
        '  </p>\n'
        '</div>'
    )
    code1.style = 'No Spacing'
    for run in code1.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("2. Add Risk Disclaimer on Investment Pages (fund-detail.html, basket-details.html):")

    code2 = doc.add_paragraph(
        '<div class="risk-disclaimer" style="background: #FEF3C7; padding: 12px; border-radius: 8px;">\n'
        '  <p style="font-size: 12px; color: #92400E;">\n'
        '    <strong>⚠️ Investment Risks:</strong> Mutual Funds are subject to market risks. '
        'Read all scheme related documents carefully before investing. '
        'Past performance is not indicative of future returns.\n'
        '  </p>\n'
        '</div>'
    )
    code2.style = 'No Spacing'
    for run in code2.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('Action Items (MANDATORY)', 2)

    actions = [
        "□ Obtain SEBI registration certificate (if not already registered)",
        "□ Add SEBI registration number to app footer on all 39 pages",
        "□ Add risk disclaimers before all 'Invest Now' buttons",
        "□ Update Terms of Service with SEBI compliance language",
        "□ Prepare documentation for App Store Connect submission notes",
        "□ Add grievance redressal contact information"
    ]

    for action in actions:
        p = doc.add_paragraph(action)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_critical_section_privacy(doc):
    """Section 3: Privacy Manifest (Required May 2024)"""
    doc.add_heading('SECTION 3: PRIVACY MANIFEST', 0)

    doc.add_heading('Rejection Risk: CRITICAL (Required Since May 1, 2024)', 1)
    risk = doc.add_paragraph()
    risk_run = risk.add_run("⛔ AUTOMATIC REJECTION IF PRIVACY MANIFEST IS MISSING")
    risk_run.font.bold = True
    risk_run.font.size = Pt(12)
    risk_run.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    doc.add_heading('Apple Requirement (As of May 2024)', 2)

    doc.add_paragraph(
        "All apps must include a PrivacyInfo.xcprivacy file that declares:\n"
        "• What data the app collects\n"
        "• Why the data is collected\n"
        "• What third-party SDKs are used\n"
        "• Required Reason API usage"
    )

    doc.add_heading('Current Status: NiveshPe App', 2)

    status = doc.add_paragraph()
    status.add_run("❌ STATUS: NOT IMPLEMENTED\n\n").font.color.rgb = COLORS['critical']
    status.add_run(
        "Your app is HTML-based, so you'll need to implement this in whichever wrapper "
        "you use (Capacitor, Cordova, or custom WKWebView)."
    )

    doc.add_heading('Data Collection in NiveshPe App (From Code Analysis)', 2)

    data_table = doc.add_table(rows=8, cols=3)
    data_table.style = 'Light Grid Accent 1'

    data_headers = data_table.rows[0].cells
    data_headers[0].text = "Data Type"
    data_headers[1].text = "Where Collected"
    data_headers[2].text = "Purpose"
    for cell in data_headers:
        cell.paragraphs[0].runs[0].font.bold = True

    data_types = [
        ("Phone Number", "auth.html (OTP)", "Account Authentication"),
        ("Name", "investor-detail.html", "KYC Compliance"),
        ("PAN Card", "investor-detail.html", "Tax & Identity Verification"),
        ("Income", "investor-detail.html", "Suitability Assessment"),
        ("Bank Details", "Add Bank flow", "Fund Transfers"),
        ("Investment Goals", "goal-create.html", "Personalization"),
        ("Device Info", "Browser/App", "Security & Fraud Prevention"),
    ]

    for i, (data_type, location, purpose) in enumerate(data_types, 1):
        row = data_table.rows[i].cells
        row[0].text = data_type
        row[1].text = location
        row[2].text = purpose

    doc.add_paragraph()

    doc.add_heading('How to Implement Privacy Manifest', 2)

    doc.add_paragraph("For Capacitor/Cordova Apps:")
    steps = [
        "1. Create PrivacyInfo.xcprivacy file in iOS project folder",
        "2. Add to Xcode project (File > Add Files to...)",
        "3. Declare all data types collected above",
        "4. List purpose for each data type",
        "5. Include third-party SDK disclosures (if using analytics, payment SDKs)"
    ]

    for step in steps:
        p = doc.add_paragraph(step)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()
    doc.add_paragraph("Sample PrivacyInfo.xcprivacy structure:")

    code = doc.add_paragraph(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE plist PUBLIC "-//Apple//DTD PLIST 1.0//EN" "...">\n'
        '<plist version="1.0">\n'
        '<dict>\n'
        '  <key>NSPrivacyCollectedDataTypes</key>\n'
        '  <array>\n'
        '    <dict>\n'
        '      <key>NSPrivacyCollectedDataType</key>\n'
        '      <string>NSPrivacyCollectedDataTypePhoneNumber</string>\n'
        '      <key>NSPrivacyCollectedDataTypeLinked</key>\n'
        '      <true/>\n'
        '      <key>NSPrivacyCollectedDataTypePurposes</key>\n'
        '      <array>\n'
        '        <string>NSPrivacyCollectedDataTypePurposeAppFunctionality</string>\n'
        '      </array>\n'
        '    </dict>\n'
        '    <!-- Add more data types as needed -->\n'
        '  </array>\n'
        '</dict>\n'
        '</plist>'
    )
    code.style = 'No Spacing'
    for run in code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    doc.add_heading('Action Items (MANDATORY)', 2)

    actions = [
        "□ Create PrivacyInfo.xcprivacy file",
        "□ Declare phone number, PAN, bank details, income, goals collection",
        "□ List purposes for each data type",
        "□ Include any third-party SDK disclosures (Google Analytics, Firebase, etc.)",
        "□ Test submission to verify manifest is recognized"
    ]

    for action in actions:
        p = doc.add_paragraph(action)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_critical_section_native_design(doc):
    """Section 3.5: Native iOS Design Patterns (NEW - CRITICAL)"""
    doc.add_heading('SECTION 3.5: NATIVE iOS DESIGN PATTERNS', 0)

    doc.add_heading('Rejection Risk: HIGH (Apple Rejects "Android-Looking" Apps)', 1)
    risk = doc.add_paragraph()
    risk_run = risk.add_run("⚠️ APPS MUST FOLLOW iOS DESIGN PATTERNS - NOT MATERIAL DESIGN")
    risk_run.font.bold = True
    risk_run.font.size = Pt(12)
    risk_run.font.color.rgb = COLORS['high']

    doc.add_paragraph()

    doc.add_heading('Apple Human Interface Guidelines (HIG) 2024 Update', 2)

    doc.add_paragraph(
        "Updated June 10, 2024: Apps that don't follow iOS design patterns risk rejection. "
        "Key quote from industry analysis: 'If you don't follow the iOS app UI/UX design guidelines, "
        "you may run into the risk of rejecting your app on the App Store.'\n\n"
        "Apple expects apps to look and feel like iOS apps, not web apps or Android ports."
    )

    doc.add_heading('Critical Native Design Violations in NiveshPe App', 2)

    violations_table = doc.add_table(rows=5, cols=3)
    violations_table.style = 'Light Grid Accent 1'

    headers = violations_table.rows[0].cells
    headers[0].text = "Component"
    headers[1].text = "Current (Wrong)"
    headers[2].text = "iOS Native (Required)"
    for cell in headers:
        cell.paragraphs[0].runs[0].font.bold = True

    violations = [
        ("Back Button", "Custom Material Symbols arrow_back icon", "iOS UINavigationBar with chevron <"),
        ("Status Bar", "Fake div with 9:41 📶🔋 emojis", "Native iOS status bar (WKWebView)"),
        ("Icons", "Material Symbols (Google/Android)", "SF Symbols (Apple) or custom iOS-style"),
        ("Design System", "Material Design 3 tokens", "iOS HIG-compliant design"),
    ]

    for i, (component, current, required) in enumerate(violations, 1):
        row = violations_table.rows[i].cells
        row[0].text = component
        row[1].text = current
        row[1].paragraphs[0].runs[0].font.color.rgb = COLORS['critical']
        row[2].text = required
        row[2].paragraphs[0].runs[0].font.color.rgb = COLORS['pass']

    doc.add_paragraph()

    doc.add_heading('Why This Matters for Rejection', 2)

    reasons = [
        "Apps that look like 'Android ports' get flagged during review",
        "Material Design is Google's design language - Apple expects iOS conventions",
        "Custom status bars signal 'web app' to reviewers (negative perception)",
        "Non-native navigation confuses iOS users familiar with standard patterns",
        "Apple prioritizes apps that 'feel at home' on iOS platform"
    ]

    for reason in reasons:
        doc.add_paragraph(f"• {reason}", style='List Bullet')

    doc.add_heading('Specific Fixes Required for NiveshPe', 2)

    doc.add_paragraph()
    doc.add_paragraph("1. REMOVE Custom Status Bar, Use Native iOS Status Bar:")

    code1 = doc.add_paragraph(
        "<!-- REMOVE THIS from ALL 39 HTML files: -->\n"
        '<div class="status-bar">\n'
        '  <div class="status-left">9:41</div>\n'
        '  <div class="status-right">📶 🔋</div>\n'
        '</div>\n\n'
        "/* Instead: Configure your iOS wrapper to show native status bar */\n"
        "// In Capacitor: capacitor.config.ts\n"
        "ios: {\n"
        "  contentInset: 'always',  // Respects safe area\n"
        "  statusBarStyle: 'dark'   // or 'light' based on background\n"
        "}"
    )
    code1.style = 'No Spacing'
    for run in code1.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("2. REPLACE Material Symbols with SF Symbols (Apple's Icon Library):")

    code2 = doc.add_paragraph(
        "<!-- Current (Wrong - Material Symbols): -->\n"
        '<span class="material-symbols-outlined">arrow_back</span>\n'
        '<span class="material-symbols-outlined">home</span>\n'
        '<span class="material-symbols-outlined">trending_up</span>\n\n'
        "<!-- iOS Native (Correct - SF Symbols via img or CSS): -->\n"
        '<img src=\"sf-symbols/chevron.left.svg\" alt=\"Back\">\n'
        '<img src=\"sf-symbols/house.svg\" alt=\"Home\">\n'
        '<img src=\"sf-symbols/chart.line.uptrend.xyaxis.svg\" alt=\"Trending\">\n\n'
        "/* Or use iOS system font icons if available in your framework */\n"
        ".ios-icon::before {\n"
        "  font-family: -apple-system, 'SF Pro Icons';\n"
        "}"
    )
    code2.style = 'No Spacing'
    for run in code2.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("3. USE iOS-Style Navigation with Native Back Button:")

    code3 = doc.add_paragraph(
        "<!-- Instead of custom back button in every page: -->\n"
        '<button class="back-btn" onclick="window.history.back()">\n'
        '  <span class="material-symbols-outlined">arrow_back</span>\n'
        '</button>\n\n'
        "<!-- Use native iOS navigation bar (implemented in iOS wrapper): -->\n"
        "// iOS (Swift/UIKit)\n"
        "navigationController?.navigationBar.prefersLargeTitles = false\n"
        "navigationItem.leftBarButtonItem = UIBarButtonItem(\n"
        "  title: \"Back\",\n"
        "  style: .plain,\n"
        "  target: self,\n"
        "  action: #selector(goBack)\n"
        ")\n\n"
        "// Or in WKWebView, intercept back navigation and use native controls"
    )
    code3.style = 'No Spacing'
    for run in code3.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading('Design System Migration Strategy', 2)

    strategy = doc.add_paragraph()
    strategy.add_run("Option 1: Full Native iOS Rewrite (Recommended for App Store success)\n").font.bold = True
    strategy.add_run(
        "• Rebuild UI using SwiftUI or UIKit\n"
        "• Automatic compliance with all iOS patterns\n"
        "• Best performance and user experience\n"
        "• ~6-8 weeks development time\n\n"
    )

    strategy.add_run("Option 2: Hybrid with Native Navigation (Faster, Good Compromise)\n").font.bold = True
    strategy.add_run(
        "• Keep HTML/CSS/JS for content\n"
        "• Use native iOS navigation bar (UINavigationController)\n"
        "• Replace status bar with native version\n"
        "• Swap Material icons for SF Symbols\n"
        "• ~2-3 weeks development time\n\n"
    )

    strategy.add_run("Option 3: Minimal Changes (Risky, May Still Get Rejected)\n").font.bold = True
    strategy.add_run(
        "• Remove custom status bar\n"
        "• Replace Material Symbols with iOS-style SVGs\n"
        "• Keep Material Design layout but adjust colors/typography\n"
        "• ~1 week development time\n"
        "• ⚠️ Still may get rejected for 'web app appearance'\n"
    )

    doc.add_paragraph()

    doc.add_heading('iOS Design Patterns Checklist', 2)

    checklist = [
        "□ Navigation bar at top (not custom header)",
        "□ iOS-style back button (< chevron, not arrow_back)",
        "□ SF Symbols icons (not Material Symbols)",
        "□ Native status bar (not custom div)",
        "□ iOS-style tab bar if using bottom navigation",
        "□ Native alerts/modals (not custom overlays)",
        "□ iOS date picker if selecting dates (spinning wheel)",
        "□ iOS-style switches (not Material Design toggles)",
        "□ Pull-to-refresh using iOS pattern",
        "□ System fonts (-apple-system, SF Pro) primary, Inter as fallback"
    ]

    for item in checklist:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading('Action Items (HIGH PRIORITY)', 2)

    actions = [
        "□ Decide on migration strategy (Full Native, Hybrid, or Minimal)",
        "□ Remove custom status bar from all 39 HTML files",
        "□ Replace Material Symbols with SF Symbols or iOS-style icons",
        "□ Implement native iOS navigation bar (if using Hybrid approach)",
        "□ Download SF Symbols app from Apple (free) to browse icons",
        "□ Test on device to ensure native elements render correctly",
        "□ Review HIG 2024 guidelines: https://developer.apple.com/design/human-interface-guidelines/"
    ]

    for action in actions:
        p = doc.add_paragraph(action)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading('SF Symbols Resources', 2)

    doc.add_paragraph(
        "SF Symbols App (free download from Apple):\n"
        "https://developer.apple.com/sf-symbols/\n\n"
        "SF Symbols provide over 5,000 configurable icons designed to integrate seamlessly with iOS. "
        "They automatically adapt to different weights, scales, and Dynamic Type settings.\n\n"
        "Common replacements for your app:\n"
        "• arrow_back → chevron.left\n"
        "• home → house.fill\n"
        "• trending_up → chart.line.uptrend.xyaxis\n"
        "• account_circle → person.circle.fill\n"
        "• search → magnifyingglass\n"
        "• filter_list → line.3.horizontal.decrease.circle"
    )

    doc.add_page_break()

def add_high_priority_ui_quality(doc):
    """Section 4: UI Quality & User Experience"""
    doc.add_heading('SECTION 4: UI QUALITY & USER EXPERIENCE', 0)

    doc.add_heading('Rejection Risk: HIGH (Subjective, But Common)', 1)
    risk = doc.add_paragraph()
    risk_run = risk.add_run("⚠️ POOR UI IS A TOP REJECTION REASON - APPLE REVIEWERS ARE STRICT")
    risk_run.font.bold = True
    risk_run.font.size = Pt(12)
    risk_run.font.color.rgb = COLORS['high']

    doc.add_paragraph()

    doc.add_heading('What Apple Looks For', 2)

    doc.add_paragraph("Apple reviewers evaluate:")

    criteria = [
        "Professional, polished design (not amateur or low-quality)",
        "Consistent navigation patterns throughout the app",
        "Clear, intuitive user flows (no confusing menus or dead ends)",
        "High-quality graphics and icons (no pixelated images)",
        "Readable text with good contrast",
        "Responsive design that works on all iPhone screen sizes"
    ]

    for criterion in criteria:
        doc.add_paragraph(f"• {criterion}", style='List Bullet')

    doc.add_heading('Current Status: NiveshPe App UI Assessment', 2)

    status = doc.add_paragraph()
    status.add_run("✓ STATUS: GOOD OVERALL QUALITY\n\n").font.color.rgb = COLORS['pass']
    status.add_run(
        "The app uses Material Design 3 with professional color scheme (blue #2563EB, green #10B981). "
        "Layout is clean and touch-optimized at 380px width."
    )

    doc.add_paragraph()

    doc.add_heading('Potential UI Concerns to Address', 2)

    concerns_table = doc.add_table(rows=6, cols=3)
    concerns_table.style = 'Light List Accent 1'

    concern_headers = concerns_table.rows[0].cells
    concern_headers[0].text = "UI Element"
    concern_headers[1].text = "Potential Issue"
    concern_headers[2].text = "Fix Needed"
    for cell in concern_headers:
        cell.paragraphs[0].runs[0].font.bold = True

    concerns = [
        ("Status Bar", "Uses emoji icons (📶🔋)", "Replace with native iOS status bar"),
        ("Back Buttons", "May be too small (40×40px)", "Test on device, ensure easy tapping"),
        ("Error States", "Not visible in all flows", "Add error messages for failed actions"),
        ("Loading States", "Some pages lack loading indicators", "Add spinners for async operations"),
        ("Empty States", "Not designed for zero data", "Add empty state designs"),
    ]

    for i, (element, issue, fix) in enumerate(concerns, 1):
        row = concerns_table.rows[i].cells
        row[0].text = element
        row[1].text = issue
        row[2].text = fix

    doc.add_paragraph()

    doc.add_heading('Specific Fixes Recommended', 2)

    doc.add_paragraph("1. Replace Emoji Status Bar with Native iOS Status Bar:")

    code1 = doc.add_paragraph(
        "<!-- REMOVE THIS from all pages: -->\n"
        '<div class="status-bar">\n'
        '  <div class="status-left">9:41</div>\n'
        '  <div class="status-right">📶 🔋</div>\n'
        '</div>\n\n'
        "<!-- USE native iOS status bar instead (configure in Capacitor/Cordova) -->\n"
        "// capacitor.config.ts\n"
        "export default {\n"
        "  ios: {\n"
        "    contentInset: 'always' // Respects safe area\n"
        "  }\n"
        "}"
    )
    code1.style = 'No Spacing'
    for run in code1.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("2. Add Proper Error States for Failed Actions:")

    code2 = doc.add_paragraph(
        "// Example for fund-detail.html investment button\n"
        '<button class="invest-btn" onclick="handleInvestment()">\n'
        '  Invest Now\n'
        '</button>\n\n'
        '<div class="error-message" id="investError" style="display:none; color: #DC2626;">\n'
        '  ⚠️ Investment failed. Please try again or contact support.\n'
        '</div>\n\n'
        "<script>\n"
        "async function handleInvestment() {\n"
        "  try {\n"
        "    // API call\n"
        "  } catch (error) {\n"
        "    document.getElementById('investError').style.display = 'block';\n"
        "  }\n"
        "}\n"
        "</script>"
    )
    code2.style = 'No Spacing'
    for run in code2.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('Action Items', 2)

    actions = [
        "□ Remove emoji status bar, use native iOS status bar",
        "□ Test all buttons on actual devices for tap responsiveness",
        "□ Add error states for all user actions (invest, add bank, verify OTP)",
        "□ Add loading spinners for async operations",
        "□ Test on different screen sizes (iPhone SE, iPhone 15 Pro Max)",
        "□ Ensure all images are high-resolution (@2x, @3x for Retina)"
    ]

    for action in actions:
        p = doc.add_paragraph(action)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_medium_priority_completeness(doc):
    """Section 5: App Completeness"""
    doc.add_heading('SECTION 5: APP COMPLETENESS', 0)

    doc.add_heading('Rejection Risk: MEDIUM-HIGH (40% of Rejections)', 1)
    risk = doc.add_paragraph()
    risk_run = risk.add_run("⚠️ INCOMPLETE APPS OR PLACEHOLDER CONTENT CAUSE 40% OF REJECTIONS")
    risk_run.font.bold = True
    risk_run.font.size = Pt(12)
    risk_run.font.color.rgb = COLORS['high']

    doc.add_paragraph()

    doc.add_heading('Apple Guideline 2.1: App Completeness', 2)

    doc.add_paragraph(
        "Submissions must be complete and ready for review. Do not submit:\n"
        "• Apps with placeholder text or 'Lorem ipsum'\n"
        "• Broken links or non-functional buttons\n"
        "• Empty states without proper designs\n"
        "• Beta versions or incomplete features\n"
        "• Apps requiring external testing or configuration"
    )

    doc.add_heading('Completeness Checklist for NiveshPe', 2)

    checklist = [
        ("✓", "All 39 pages accessible and functional", COLORS['pass']),
        ("⚠️", "Terms & Privacy Policy links - verify they work", COLORS['medium']),
        ("⚠️", "Customer support contact - must be responsive", COLORS['medium']),
        ("⚠️", "Demo account for App Review team - required for testing", COLORS['medium']),
        ("✓", "No Lorem Ipsum text found", COLORS['pass']),
        ("⚠️", "API endpoints - must be production-ready, not staging", COLORS['medium']),
    ]

    for symbol, item, color in checklist:
        p = doc.add_paragraph(f"{symbol} {item}")
        p.paragraph_format.left_indent = Inches(0.25)
        if color != COLORS['pass']:
            p.runs[0].font.color.rgb = color

    doc.add_heading('Critical: Demo Account for App Review', 2)

    demo = doc.add_paragraph()
    demo.add_run("MANDATORY: ").font.bold = True
    demo.add_run(
        "You must provide a demo account with pre-loaded data so Apple reviewers can "
        "test the app without creating real investments.\n\n"
    )
    demo.add_run("Demo Account Requirements:\n").font.bold = True

    demo_reqs = [
        "• Phone number that can receive OTP (or bypass OTP for this account)",
        "• Pre-filled KYC details",
        "• Sample portfolio with 3-5 funds showing returns",
        "• Functional 'Invest Now' flow (can use test payment gateway)",
        "• Goals and basket recommendations populated"
    ]

    for req in demo_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_paragraph()

    note = doc.add_paragraph()
    note.add_run("Note: ").font.bold = True
    note.add_run(
        "Include demo credentials in App Store Connect > App Review Information > Notes field:\n"
        "Phone: +91-XXXXXXXXXX\n"
        "OTP: 123456 (or auto-verify)\n"
        "Test Mode: Investment flows use sandbox payment gateway"
    )

    doc.add_heading('Action Items', 2)

    actions = [
        "□ Test ALL links (Terms, Privacy, Support, SEBI disclosures)",
        "□ Verify no placeholder text remains",
        "□ Create dedicated demo account with sample data",
        "□ Configure OTP bypass or test phone number",
        "□ Set up test payment gateway for investment flows",
        "□ Add demo credentials to App Store Connect submission notes",
        "□ Verify customer support email/phone are monitored"
    ]

    for action in actions:
        p = doc.add_paragraph(action)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_testing_checklist(doc):
    """Section 6: Pre-Submission Testing Checklist"""
    doc.add_heading('SECTION 6: PRE-SUBMISSION TESTING CHECKLIST', 0)

    doc.add_heading('Complete Before Submitting to App Store', 1)

    doc.add_heading('Device Testing (CRITICAL)', 2)

    device_tests = [
        "□ Test on iPhone SE (small screen)",
        "□ Test on iPhone 14/15 (standard size)",
        "□ Test on iPhone 15 Pro Max (large screen)",
        "□ Test on iOS 17.0, 17.5, and 18.0+",
        "□ No crashes on any device/version combination",
        "□ UI looks correct on all screen sizes",
        "□ Touch targets are easy to tap on all devices"
    ]

    for test in device_tests:
        p = doc.add_paragraph(test)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading('Functionality Testing', 2)

    func_tests = [
        "□ OTP verification works",
        "□ KYC form submission completes",
        "□ Fund search and filtering functional",
        "□ Investment flow (SIP/Lumpsum) completes",
        "□ Portfolio displays correctly",
        "□ Goals creation and tracking works",
        "□ Back buttons navigate correctly",
        "□ All external links open (Terms, Privacy, Support)"
    ]

    for test in func_tests:
        p = doc.add_paragraph(test)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading('Compliance Verification', 2)

    compliance_tests = [
        "□ SEBI registration number visible on all pages",
        "□ Risk disclaimers shown before investments",
        "□ Privacy manifest (PrivacyInfo.xcprivacy) included",
        "□ Native iOS design patterns (SF Symbols, native navigation)",
        "□ Custom status bar removed, using native iOS status bar",
        "□ SEBI documentation uploaded to App Store Connect",
        "□ Demo account credentials provided in review notes",
        "□ Terms of Service and Privacy Policy links work"
    ]

    for test in compliance_tests:
        p = doc.add_paragraph(test)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading('Final Quality Checks', 2)

    quality_tests = [
        "□ No placeholder content (Lorem ipsum, TODO, etc.)",
        "□ All images are high-resolution",
        "□ No broken images or missing icons",
        "□ Error states designed for all user actions",
        "□ Loading states for async operations",
        "□ Native iOS status bar (no emoji)",
        "□ App metadata (name, description, screenshots) ready"
    ]

    for test in quality_tests:
        p = doc.add_paragraph(test)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_appendix(doc):
    """Appendix with references"""
    doc.add_heading('APPENDIX: REFERENCES & RESOURCES', 0)

    doc.add_heading('Apple Guidelines Referenced', 2)

    refs = [
        ("Guideline 2.1", "App Completeness", "https://developer.apple.com/app-store/review/guidelines/#app-completeness"),
        ("Guideline 3.2.1(viii)", "Financial Services", "https://developer.apple.com/app-store/review/guidelines/#financial-services"),
        ("Guideline 4.0", "Design", "https://developer.apple.com/design/human-interface-guidelines/"),
        ("Guideline 5.1", "Privacy", "https://developer.apple.com/app-store/review/guidelines/#privacy"),
        ("Privacy Manifest", "Required since May 2024", "https://developer.apple.com/documentation/bundleresources/privacy_manifest_files"),
    ]

    for guideline, description, url in refs:
        p = doc.add_paragraph()
        p.add_run(f"{guideline}: {description}\n").font.bold = True
        p.add_run(url).font.color.rgb = COLORS['blue']

    doc.add_paragraph()

    doc.add_heading('SEBI Regulations Referenced', 2)

    sebi_refs = [
        "SEBI (Mutual Funds) Regulations, 1996",
        "SEBI (Investment Advisers) Regulations, 2013",
        "SEBI Circular on Risk-o-Meter (December 2024)"
    ]

    for ref in sebi_refs:
        doc.add_paragraph(f"• {ref}", style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('2024 Rejection Statistics (Sources)', 2)

    doc.add_paragraph(
        "• Apple reviewed 7.77 million submissions in 2024\n"
        "• 24.8% rejection rate (approximately 1.93 million rejections)\n"
        "• 40% of rejections related to Guideline 2.1 (Completeness)\n"
        "• Fintech apps face additional scrutiny per Guideline 3.2.1\n\n"
        "Sources:\n"
        "- UXCam: Top 10 App Store Rejection Reasons (2024)\n"
        "- MobileLoud: 16 Reasons Why Your App Could Be Rejected\n"
        "- Apple Developer: App Store Review Guidelines"
    )

    doc.add_paragraph()

    doc.add_heading('Recommended Tools', 2)

    tools = [
        "Safari Web Inspector - Debug iOS web apps",
        "Xcode - Required for Privacy Manifest creation",
        "TestFlight - Beta testing before submission",
        "App Store Connect - Submission and review management",
        "Capacitor/Cordova - Web-to-native app frameworks"
    ]

    for tool in tools:
        doc.add_paragraph(f"• {tool}", style='List Bullet')

def save_document(doc, filename):
    """Save the document"""
    doc.save(filename)
    file_size = os.path.getsize(filename) / (1024 * 1024)  # Convert to MB

    print("\n" + "="*60)
    print("✓ Document saved:", filename)
    print(f"✓ File size: {file_size:.2f} MB")
    print(f"✓ Location: {os.path.abspath(filename)}")
    print("="*60)
    print("\nDocument generation complete!")
    print("\nNext steps:")
    print("1. Open the document in Microsoft Word")
    print("2. Complete all checklist items marked with □")
    print("3. Gather SEBI documentation for App Store Connect")
    print("4. Test on actual iOS devices before submission")

def main():
    """Main execution function"""
    print("\nGenerating iOS App Store ACTUAL Rejection Risks Document...")
    print("="*60)

    doc = create_document()
    print("✓ Document created")

    add_title_page(doc)
    print("✓ Title page added")

    add_executive_summary(doc)
    print("✓ Executive summary added")

    add_critical_section_crashes(doc)
    print("✓ Section 1: Crashes & Stability added")

    add_critical_section_financial(doc)
    print("✓ Section 2: Financial Compliance added")

    add_critical_section_privacy(doc)
    print("✓ Section 3: Privacy Manifest added")

    add_critical_section_native_design(doc)
    print("✓ Section 3.5: Native iOS Design Patterns added")

    add_high_priority_ui_quality(doc)
    print("✓ Section 4: UI Quality added")

    add_medium_priority_completeness(doc)
    print("✓ Section 5: App Completeness added")

    add_testing_checklist(doc)
    print("✓ Section 6: Testing Checklist added")

    add_appendix(doc)
    print("✓ Appendix added")

    filename = "iOS_AppStore_ACTUAL_Rejection_Risks_NiveshPe.docx"
    save_document(doc, filename)

if __name__ == "__main__":
    main()
