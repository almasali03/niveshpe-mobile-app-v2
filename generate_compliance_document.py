#!/usr/bin/env python3
"""
iOS App Store Design Guidelines Compliance Document Generator
Creates a comprehensive Microsoft Word document detailing rejection-risk design violations
for the NiveshPe mutual fund investment app.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Color definitions matching the plan
COLORS = {
    'critical': RGBColor(239, 68, 68),    # #EF4444
    'high': RGBColor(245, 158, 11),       # #F59E0B
    'medium': RGBColor(251, 191, 36),     # #FBBF24
    'pass': RGBColor(16, 185, 129),       # #10B981
    'blue': RGBColor(37, 99, 235),        # #2563EB
}

def create_document():
    """Create and configure the Word document"""
    doc = Document()

    # Set document properties
    doc.core_properties.title = "iOS App Store Design Compliance Audit - NiveshPe App"
    doc.core_properties.author = "NiveshPe Development Team"
    doc.core_properties.comments = "Comprehensive audit of iOS App Store design guideline violations with rejection risk assessment"

    return doc

def add_title_page(doc):
    """Add formatted title page"""
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("iOS App Store\nDesign Compliance Audit")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = COLORS['blue']
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("NiveshPe Mutual Fund Investment App")
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run("December 2024")
    date_run.font.size = Pt(14)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Document info box
    info = doc.add_paragraph()
    info_run = info.add_run(
        "Document Type: Design Guidelines Compliance Audit\n"
        "Priority Focus: App Store Rejection Risks\n"
        "Pages Analyzed: 39 HTML files\n"
        "Compliance Score: 45/100 (FAIL - Not Ready for Submission)"
    )
    info_run.font.size = Pt(11)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page break
    doc.add_page_break()

def add_executive_summary(doc):
    """Add executive summary section"""
    doc.add_heading('EXECUTIVE SUMMARY', 0)

    # Critical Issues Overview
    doc.add_heading('Overall Compliance Assessment', 1)

    p = doc.add_paragraph()
    run = p.add_run("COMPLIANCE SCORE: 45/100 ")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLORS['critical']

    fail_run = p.add_run("(FAIL - Not Ready for Submission)")
    fail_run.font.size = Pt(14)
    fail_run.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    # Critical violations table
    doc.add_heading('Critical Violations Summary', 2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Issue'
    hdr_cells[1].text = 'Compliance'
    hdr_cells[2].text = 'Rejection Risk'
    hdr_cells[3].text = 'Priority'

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add critical violations
    violations = [
        ('VoiceOver Accessibility', '0%', 'CRITICAL', 'Week 1-2'),
        ('SEBI Financial Disclosures', '0%', 'CRITICAL', 'Week 1'),
        ('Dynamic Type Support', '0%', 'HIGH', 'Week 3-4'),
        ('Touch Target Sizes', '60%', 'MEDIUM-HIGH', 'Week 1-2'),
        ('Privacy Transparency', '30%', 'HIGH', 'Week 2-3'),
    ]

    for issue, compliance, risk, priority in violations:
        row_cells = table.add_row().cells
        row_cells[0].text = issue
        row_cells[1].text = compliance
        row_cells[2].text = risk
        row_cells[3].text = priority

    doc.add_paragraph()

    # Immediate actions
    doc.add_heading('Immediate Actions Required (Before Submission)', 2)

    doc.add_paragraph(
        "The following critical issues MUST be fixed before App Store submission to avoid rejection:",
        style='List Bullet'
    )

    critical_actions = [
        "Add aria-label accessibility attributes to all ~450 interactive elements across 39 pages",
        "Display SEBI registration number in footer of all pages",
        'Add "Mutual Funds are subject to market risks" disclaimer on all fund detail pages',
        "Increase back button touch targets from 40×40px to minimum 44×44pt (36 pages affected)",
        "Associate all form input labels with their inputs using proper HTML attributes",
        "Replace emoji status bar icons (📶🔋) with accessible alternatives",
    ]

    for action in critical_actions:
        p = doc.add_paragraph(action, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Remediation effort
    doc.add_heading('Estimated Remediation Effort', 2)

    effort_table = doc.add_table(rows=1, cols=3)
    effort_table.style = 'Light List Accent 1'

    effort_hdr = effort_table.rows[0].cells
    effort_hdr[0].text = 'Priority Level'
    effort_hdr[1].text = 'Developer Hours'
    effort_hdr[2].text = 'Timeline'

    for cell in effort_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    effort_data = [
        ('Critical Fixes', '60-80 hours', 'Week 1-2'),
        ('High Priority', '50-70 hours', 'Week 3-4'),
        ('Medium Priority', '30-50 hours', 'Week 5-6'),
        ('Testing & Validation', '20-30 hours', 'Week 7'),
    ]

    for priority, hours, timeline in effort_data:
        row = effort_table.add_row().cells
        row[0].text = priority
        row[1].text = hours
        row[2].text = timeline

    doc.add_paragraph()

    total = doc.add_paragraph()
    total_run = total.add_run("TOTAL ESTIMATED EFFORT: 160-230 developer hours")
    total_run.font.bold = True
    total_run.font.size = Pt(12)

    timeline = doc.add_paragraph()
    timeline_run = timeline.add_run("RECOMMENDED TIMELINE: 6-8 weeks before App Store submission")
    timeline_run.font.bold = True
    timeline_run.font.size = Pt(12)
    timeline_run.font.color.rgb = COLORS['high']

    doc.add_page_break()

def add_accessibility_section(doc):
    """Add detailed accessibility violations section"""
    doc.add_heading('SECTION 2: CRITICAL VIOLATIONS - ACCESSIBILITY', 0)

    # 2.1 VoiceOver Support
    doc.add_heading('2.1 VoiceOver & Screen Reader Support', 1)

    guideline_box = doc.add_paragraph()
    guideline_box.add_run("Apple Guideline: ").font.bold = True
    guideline_box.add_run("4.0 (Design - Accessibility)\n")
    guideline_box.add_run("HIG Reference: ").font.bold = True
    guideline_box.add_run("https://developer.apple.com/design/human-interface-guidelines/accessibility\n")
    guideline_box.add_run("WCAG Reference: ").font.bold = True
    guideline_box.add_run("4.1.2 Name, Role, Value\n")
    guideline_box.add_run("Rejection Risk: ").font.bold = True
    risk_run = guideline_box.add_run("CRITICAL")
    risk_run.font.color.rgb = COLORS['critical']
    risk_run.font.bold = True

    doc.add_paragraph()

    # Current Status
    status = doc.add_paragraph()
    status.add_run("CURRENT STATUS: ").font.bold = True
    fail_run = status.add_run("❌ 0% COMPLIANT")
    fail_run.font.color.rgb = COLORS['critical']
    fail_run.font.bold = True

    doc.add_paragraph()

    # Audit Findings
    doc.add_heading('Audit Findings', 2)

    findings_table = doc.add_table(rows=1, cols=2)
    findings_table.style = 'Medium Shading 1 Accent 1'

    findings_hdr = findings_table.rows[0].cells
    findings_hdr[0].text = 'Metric'
    findings_hdr[1].text = 'Value'

    for cell in findings_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    findings_data = [
        ('Total Pages Analyzed', '39'),
        ('Interactive Elements Found', '~450'),
        ('Elements with Accessibility Labels', '0'),
        ('Compliance Rate', '0%'),
        ('Pages Affected', '39 (100%)'),
    ]

    for metric, value in findings_data:
        row = findings_table.add_row().cells
        row[0].text = metric
        row[1].text = value

    doc.add_paragraph()

    # Specific Violations
    doc.add_heading('Specific Violations', 2)

    # Violation 1: Back Buttons
    doc.add_heading('1. Back Buttons (All Pages)', 3)

    doc.add_paragraph(
        "Location: auth.html, fund-detail.html, basket-details.html, mutual-funds.html, etc."
    )

    doc.add_paragraph("Current Code:")
    current_code = doc.add_paragraph(
        '<button class=\"back-btn\" onclick=\"history.back()\">\n'
        '    <span class=\"material-symbols-outlined\">arrow_back</span>\n'
        '</button>',
        style='Intense Quote'
    )
    current_code.paragraph_format.left_indent = Inches(0.5)
    for run in current_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()

    issue = doc.add_paragraph()
    issue.add_run("VoiceOver Reads: ").font.bold = True
    issue_run = issue.add_run('"Button" (no description)')
    issue_run.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    doc.add_paragraph("Required Fix:")
    fixed_code = doc.add_paragraph(
        '<button class=\"back-btn\" \n'
        '        onclick=\"history.back()\"\n'
        '        aria-label=\"Go back to previous page\">\n'
        '    <span class=\"material-symbols-outlined\" aria-hidden=\"true\">arrow_back</span>\n'
        '</button>',
        style='Intense Quote'
    )
    fixed_code.paragraph_format.left_indent = Inches(0.5)
    for run in fixed_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['pass']

    doc.add_paragraph()

    pages_para = doc.add_paragraph()
    pages_para.add_run("Pages Affected: ").font.bold = True
    pages_para.add_run("36\n")
    pages_para.add_run("Priority: ").font.bold = True
    priority_run = pages_para.add_run("CRITICAL")
    priority_run.font.color.rgb = COLORS['critical']
    priority_run.font.bold = True

    doc.add_paragraph()

    # Violation 2: Status Bar Icons
    doc.add_heading('2. Status Bar Icons (All Pages)', 3)

    doc.add_paragraph("Current Code:")
    status_code = doc.add_paragraph(
        '<div class=\"status-icons\">\n'
        '    <span>📶</span>\n'
        '    <span>📡</span>\n'
        '    <span>🔋</span>\n'
        '</div>',
        style='Intense Quote'
    )
    status_code.paragraph_format.left_indent = Inches(0.5)
    for run in status_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()

    issue2 = doc.add_paragraph()
    issue2.add_run("Issue: ").font.bold = True
    issue2.add_run("Emoji read literally by VoiceOver, not meaningful. VoiceOver announces \"Signal strength emoji\" instead of actual signal level.")

    doc.add_paragraph()

    doc.add_paragraph("Required Fix:")
    fixed_status = doc.add_paragraph(
        '<div class=\"status-icons\" aria-label=\"System status\">\n'
        '    <span class=\"signal-icon\" aria-label=\"Signal: 4 bars\">📶</span>\n'
        '    <span class=\"wifi-icon\" aria-label=\"WiFi connected\">📡</span>\n'
        '    <span class=\"battery-icon\" aria-label=\"Battery: 85%\">🔋</span>\n'
        '</div>',
        style='Intense Quote'
    )
    fixed_status.paragraph_format.left_indent = Inches(0.5)
    for run in fixed_status.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['pass']

    doc.add_paragraph()

    pages_para2 = doc.add_paragraph()
    pages_para2.add_run("Pages Affected: ").font.bold = True
    pages_para2.add_run("39\n")
    pages_para2.add_run("Priority: ").font.bold = True
    priority_run2 = pages_para2.add_run("HIGH")
    priority_run2.font.color.rgb = COLORS['high']
    priority_run2.font.bold = True

    doc.add_paragraph()

    # Violation 3: Investment Buttons
    doc.add_heading('3. Investment CTA Buttons (fund-detail.html)', 3)

    doc.add_paragraph("Current Code:")
    cta_code = doc.add_paragraph(
        '<button class=\"invest-cta\">\n'
        '    Invest Now\n'
        '</button>',
        style='Intense Quote'
    )
    cta_code.paragraph_format.left_indent = Inches(0.5)
    for run in cta_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()

    issue3 = doc.add_paragraph()
    issue3.add_run("Issue: ").font.bold = True
    issue3.add_run("Button lacks context - users don't know what fund they're investing in when navigating with VoiceOver.")

    doc.add_paragraph()

    doc.add_paragraph("Required Fix:")
    fixed_cta = doc.add_paragraph(
        '<button class=\"invest-cta\"\n'
        '        aria-label=\"Invest in HDFC Large Cap Fund - Review fund details before investing\">\n'
        '    Invest Now\n'
        '</button>',
        style='Intense Quote'
    )
    fixed_cta.paragraph_format.left_indent = Inches(0.5)
    for run in fixed_cta.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['pass']

    doc.add_paragraph()

    pages_para3 = doc.add_paragraph()
    pages_para3.add_run("Pages Affected: ").font.bold = True
    pages_para3.add_run("fund-detail.html, basket-details.html, mutual-funds.html\n")
    pages_para3.add_run("Priority: ").font.bold = True
    priority_run3 = pages_para3.add_run("CRITICAL")
    priority_run3.font.color.rgb = COLORS['critical']
    priority_run3.font.bold = True

    doc.add_paragraph()

    # Action Items
    doc.add_heading('Action Items', 2)

    actions = [
        "Add aria-label to all back buttons across 36 pages (Timeline: Week 1, Effort: 12 hours)",
        "Add meaningful aria-labels to all Material Symbols icons (Timeline: Week 1, Effort: 8 hours)",
        "Associate all form input labels with inputs using for/id attributes (Timeline: Week 1-2, Effort: 6 hours)",
        "Add aria-labels to all status bar elements (Timeline: Week 1, Effort: 2 hours)",
        "Test with VoiceOver on actual iOS device across all 39 pages (Timeline: Week 2, Effort: 8 hours)",
    ]

    for i, action in enumerate(actions, 1):
        p = doc.add_paragraph(f"☐ ACTION {i}: {action}", style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Testing Procedure
    doc.add_heading('Testing Procedure', 2)

    doc.add_paragraph(
        "1. Enable VoiceOver on iOS device (Settings > Accessibility > VoiceOver)",
        style='List Number'
    )
    doc.add_paragraph(
        "2. Navigate through each page using swipe gestures only (no visual reference)",
        style='List Number'
    )
    doc.add_paragraph(
        "3. Verify every button/link announces its purpose clearly",
        style='List Number'
    )
    doc.add_paragraph(
        "4. Test complete investment flow hands-free using only VoiceOver",
        style='List Number'
    )
    doc.add_paragraph(
        "5. Use Xcode Accessibility Inspector for automated checks",
        style='List Number'
    )
    doc.add_paragraph(
        "6. Document any elements that still lack proper labels",
        style='List Number'
    )

    doc.add_page_break()

    # 2.2 Touch Target Sizes
    doc.add_heading('2.2 Touch Target Size Requirements', 1)

    guideline_box2 = doc.add_paragraph()
    guideline_box2.add_run("Apple HIG: ").font.bold = True
    guideline_box2.add_run("44×44pt minimum touch target size\n")
    guideline_box2.add_run("Reference: ").font.bold = True
    guideline_box2.add_run("https://developer.apple.com/design/human-interface-guidelines/layout\n")
    guideline_box2.add_run("Rejection Risk: ").font.bold = True
    risk_run2 = guideline_box2.add_run("MEDIUM-HIGH")
    risk_run2.font.color.rgb = COLORS['high']
    risk_run2.font.bold = True

    doc.add_paragraph()

    # Current Status
    status2 = doc.add_paragraph()
    status2.add_run("CURRENT STATUS: ").font.bold = True
    partial_run = status2.add_run("⚠️ 60% COMPLIANT")
    partial_run.font.color.rgb = COLORS['medium']
    partial_run.font.bold = True

    doc.add_paragraph()

    doc.add_heading('Violations Found', 2)

    touch_table = doc.add_table(rows=1, cols=4)
    touch_table.style = 'Light Grid Accent 1'

    touch_hdr = touch_table.rows[0].cells
    touch_hdr[0].text = 'Element Type'
    touch_hdr[1].text = 'Current Size'
    touch_hdr[2].text = 'Required Size'
    touch_hdr[3].text = 'Pages Affected'

    for cell in touch_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    touch_data = [
        ('Back buttons', '40×40px', '44×44pt', '36'),
        ('Close buttons (modals)', '32×32px', '44×44pt', '8'),
        ('Period selector buttons', '38×38px', '44×44pt', '12'),
        ('Icon-only action buttons', '40×40px', '44×44pt', '15'),
    ]

    for element, current, required, pages in touch_data:
        row = touch_table.add_row().cells
        row[0].text = element
        row[1].text = current
        row[2].text = required
        row[3].text = pages

    doc.add_paragraph()

    # Fix Example
    doc.add_heading('Fix Example: Back Button', 2)

    doc.add_paragraph("Current CSS:")
    current_css = doc.add_paragraph(
        '.back-btn {\n'
        '    width: 40px;\n'
        '    height: 40px;\n'
        '    border-radius: 50%;\n'
        '}',
        style='Intense Quote'
    )
    current_css.paragraph_format.left_indent = Inches(0.5)
    for run in current_css.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_paragraph("Fixed CSS:")
    fixed_css = doc.add_paragraph(
        '.back-btn {\n'
        '    width: 44px;  /* Updated to meet minimum */\n'
        '    height: 44px; /* Updated to meet minimum */\n'
        '    border-radius: 50%;\n'
        '    /* Visual size can be smaller with padding if needed */\n'
        '    padding: 8px; /* Creates visual size of 28px icon */\n'
        '}',
        style='Intense Quote'
    )
    fixed_css.paragraph_format.left_indent = Inches(0.5)
    for run in fixed_css.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['pass']

    doc.add_paragraph()

    # Action Items
    doc.add_heading('Action Items', 2)

    touch_actions = [
        "Update all back buttons to 44×44pt minimum (36 pages, Week 1, 4 hours)",
        "Increase modal close buttons to 44×44pt (8 pages, Week 1, 2 hours)",
        "Fix period selector button sizes (12 pages, Week 1-2, 3 hours)",
        "Audit all icon-only buttons and increase to minimum (15 pages, Week 2, 4 hours)",
        "Test touch targets with motor disability simulator (Week 2, 2 hours)",
    ]

    for i, action in enumerate(touch_actions, 1):
        p = doc.add_paragraph(f"☐ ACTION {i}: {action}", style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

    # 2.3 Dynamic Type
    doc.add_heading('2.3 Dynamic Type & Text Scaling Support', 1)

    guideline_box3 = doc.add_paragraph()
    guideline_box3.add_run("Apple Guideline: ").font.bold = True
    guideline_box3.add_run("4.0 (Accessibility - Dynamic Type)\n")
    guideline_box3.add_run("HIG Reference: ").font.bold = True
    guideline_box3.add_run("https://developer.apple.com/design/human-interface-guidelines/foundations/typography\n")
    guideline_box3.add_run("WCAG Reference: ").font.bold = True
    guideline_box3.add_run("1.4.4 Resize Text (Level AA)\n")
    guideline_box3.add_run("Rejection Risk: ").font.bold = True
    risk_run3 = guideline_box3.add_run("HIGH")
    risk_run3.font.color.rgb = COLORS['high']
    risk_run3.font.bold = True

    doc.add_paragraph()

    # Current Status
    status3 = doc.add_paragraph()
    status3.add_run("CURRENT STATUS: ").font.bold = True
    fail_run3 = status3.add_run("❌ 0% COMPLIANT")
    fail_run3.font.color.rgb = COLORS['critical']
    fail_run3.font.bold = True

    doc.add_paragraph()

    doc.add_heading('The Problem', 2)

    doc.add_paragraph(
        "All typography in the app uses fixed pixel (px) values defined in md3-tokens.css. "
        "This means text cannot scale when users adjust their iOS Dynamic Type settings. "
        "Apple requires apps to support text scaling up to 200% minimum, with graceful degradation up to 300%."
    )

    doc.add_paragraph()

    doc.add_paragraph("Current Implementation (md3-tokens.css):")
    current_type = doc.add_paragraph(
        ':root {\n'
        '  --md-sys-typescale-display-large-size: 57px;\n'
        '  --md-sys-typescale-headline-medium-size: 28px;\n'
        '  --md-sys-typescale-body-large-size: 16px;\n'
        '  /* All values are fixed px - cannot scale */\n'
        '}',
        style='Intense Quote'
    )
    current_type.paragraph_format.left_indent = Inches(0.5)
    for run in current_type.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading('Required Fix: Migration to rem Units', 2)

    doc.add_paragraph(
        "Convert all typography from px to rem units. The base font size (16px = 1rem) "
        "will scale with user's system preferences."
    )

    doc.add_paragraph()

    doc.add_paragraph("Updated Implementation:")
    fixed_type = doc.add_paragraph(
        ':root {\n'
        '  /* Base: 16px = 1rem */\n'
        '  --md-sys-typescale-display-large-size: 3.5625rem;  /* 57px/16 */\n'
        '  --md-sys-typescale-headline-medium-size: 1.75rem;  /* 28px/16 */\n'
        '  --md-sys-typescale-body-large-size: 1rem;         /* 16px/16 */\n'
        '}\n\n'
        '/* Support iOS Dynamic Type categories */\n'
        '@media (prefers-reduced-motion: no-preference) {\n'
        '  html {\n'
        '    font-size: 16px; /* Base size */\n'
        '  }\n'
        '}\n\n'
        '/* Scale for larger accessibility text */\n'
        '@media screen and (min-width: 0) {\n'
        '  html {\n'
        '    font-size: calc(16px + 0.2vw); /* Responsive scaling */\n'
        '  }\n'
        '}',
        style='Intense Quote'
    )
    fixed_type.paragraph_format.left_indent = Inches(0.5)
    for run in fixed_type.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS['pass']

    doc.add_paragraph()

    # Migration Strategy
    doc.add_heading('Migration Strategy', 2)

    migration_steps = [
        "Audit md3-tokens.css and create px to rem conversion table (Week 3, 4 hours)",
        "Update all typography tokens to use rem units (Week 3, 8 hours)",
        "Test all 39 pages at 100%, 150%, 200% text scaling (Week 3, 12 hours)",
        "Fix container height/width issues that break with larger text (Week 4, 16 hours)",
        "Implement max-width constraints on containers to prevent overflow (Week 4, 8 hours)",
        "Add CSS clamp() for responsive scaling: clamp(1rem, 2.5vw, 1.5rem) (Week 4, 6 hours)",
        "Test with real users requiring large text accessibility (Week 4, 4 hours)",
    ]

    for i, step in enumerate(migration_steps, 1):
        p = doc.add_paragraph(f"☐ STEP {i}: {step}", style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    doc.add_heading('Testing Checklist', 2)

    test_steps = [
        "Enable Large Text in iOS Settings > Accessibility > Display & Text Size",
        "Set text size to maximum (310% of default)",
        "Navigate through all app pages and verify text is readable",
        "Ensure no text is cut off or overlapping",
        "Verify buttons and interactive elements remain functional",
        "Test with middle sizes (150%, 200%) for graceful degradation",
        "Document any pages that need layout adjustments",
    ]

    for step in test_steps:
        p = doc.add_paragraph(f"☐ {step}", style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_financial_compliance_section(doc):
    """Add SEBI financial compliance section"""
    doc.add_heading('SECTION 3: CRITICAL VIOLATIONS - FINANCIAL COMPLIANCE', 0)

    # 3.1 SEBI Requirements
    doc.add_heading('3.1 SEBI Regulatory Requirements', 1)

    guideline_box = doc.add_paragraph()
    guideline_box.add_run("SEBI Regulations: ").font.bold = True
    guideline_box.add_run("SEBI (Mutual Funds) Regulations 1996\n")
    guideline_box.add_run("Apple Guideline: ").font.bold = True
    guideline_box.add_run("3.2.1(viii) - Financial Services\n")
    guideline_box.add_run("Rejection Risk: ").font.bold = True
    risk_run = guideline_box.add_run("CRITICAL")
    risk_run.font.color.rgb = COLORS['critical']
    risk_run.font.bold = True
    risk_run2 = guideline_box.add_run(" (Both Apple & Legal)")
    risk_run2.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    # Current Status
    status = doc.add_paragraph()
    status.add_run("CURRENT STATUS: ").font.bold = True
    partial_run = status.add_run("⚠️ PARTIAL COMPLIANCE (40%)")
    partial_run.font.color.rgb = COLORS['medium']
    partial_run.font.bold = True

    doc.add_paragraph()

    doc.add_heading('Mandatory SEBI Disclosures', 2)

    # Disclosure table
    sebi_table = doc.add_table(rows=1, cols=4)
    sebi_table.style = 'Medium Shading 1 Accent 2'

    sebi_hdr = sebi_table.rows[0].cells
    sebi_hdr[0].text = 'Requirement'
    sebi_hdr[1].text = 'Status'
    sebi_hdr[2].text = 'Location Required'
    sebi_hdr[3].text = 'Priority'

    for cell in sebi_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    sebi_data = [
        ('SEBI Registration Number', '❌ Missing', 'Footer of all pages', 'CRITICAL'),
        ('Market Risks Disclaimer', '❌ Missing', 'Fund detail pages', 'CRITICAL'),
        ('Risk Categorization', '⚠️ Partial', 'All fund cards', 'HIGH'),
        ('Past Performance Disclaimer', '❌ Missing', 'Returns display', 'HIGH'),
        ('NAV Display with Date', '✓ Present', 'fund-detail.html', 'Pass'),
        ('Expense Ratio Disclosure', '✓ Present', 'Fund details', 'Pass'),
        ('Grievance Redressal', '❌ Missing', 'Footer/Help section', 'HIGH'),
    ]

    for requirement, status, location, priority in sebi_data:
        row = sebi_table.add_row().cells
        row[0].text = requirement
        row[1].text = status
        row[2].text = location
        row[3].text = priority

    doc.add_paragraph()

    # Critical Missing Items
    doc.add_heading('1. SEBI Registration Number', 3)

    req1 = doc.add_paragraph()
    req1.add_run("Required: ").font.bold = True
    req1.add_run('"Registered with SEBI as Mutual Fund Distributor (ARN-XXXXX)"\n')
    req1.add_run("Current: ").font.bold = True
    current_run = req1.add_run("Not found on any page")
    current_run.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    doc.add_paragraph("Implementation:")
    sebi_impl = doc.add_paragraph(
        '<!-- Add to footer of ALL 39 pages -->\n'
        '<footer class=\"app-footer\">\n'
        '    <div class=\"footer-content\">\n'
        '        <p class=\"sebi-registration\">\n'
        '            <strong>[Your Company Name]</strong> is registered with SEBI as a \n'
        '            Mutual Fund Distributor (Registration No: ARN-XXXXXX). \n'
        '            Valid till: DD/MM/YYYY\n'
        '        </p>\n'
        '        <p class=\"regulatory-links\">\n'
        '            <a href=\"/privacy-policy\">Privacy Policy</a> | \n'
        '            <a href=\"/terms\">Terms of Service</a> | \n'
        '            <a href=\"/grievances\">Grievance Redressal</a>\n'
        '        </p>\n'
        '    </div>\n'
        '</footer>',
        style='Intense Quote'
    )
    sebi_impl.paragraph_format.left_indent = Inches(0.5)
    for run in sebi_impl.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS['pass']

    doc.add_paragraph()

    impl1 = doc.add_paragraph()
    impl1.add_run("Location: ").font.bold = True
    impl1.add_run("Footer of all 39 pages\n")
    impl1.add_run("Priority: ").font.bold = True
    priority1 = impl1.add_run("CRITICAL - Week 1, Day 1")
    priority1.font.color.rgb = COLORS['critical']
    priority1.font.bold = True

    doc.add_paragraph()

    # Risk Disclaimer
    doc.add_heading('2. Standard Market Risk Disclaimer', 3)

    req2 = doc.add_paragraph()
    req2.add_run("Required: ").font.bold = True
    req2.add_run(
        '"Mutual Funds are subject to market risks. Please read all '
        'scheme related documents carefully before investing."\n'
    )
    req2.add_run("Current: ").font.bold = True
    current_run2 = req2.add_run("Not found on fund detail pages")
    current_run2.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    doc.add_paragraph("Implementation Example:")
    disclaimer_impl = doc.add_paragraph(
        '<!-- Add to fund-detail.html BEFORE "Invest Now" button -->\n'
        '<div class=\"sebi-disclaimer-box\">\n'
        '    <div class=\"disclaimer-header\">\n'
        '        <span class=\"material-symbols-outlined\" aria-hidden=\"true\">info</span>\n'
        '        <strong>Investment Risk Disclosure</strong>\n'
        '    </div>\n'
        '    <div class=\"risk-level moderate\">\n'
        '        <span>Risk Level: </span>\n'
        '        <strong>Moderately High</strong>\n'
        '    </div>\n'
        '    <p class=\"disclaimer-text\">\n'
        '        <strong>Mutual Funds are subject to market risks.</strong> \n'
        '        Please read all scheme related documents carefully before investing. \n'
        '        Past performance is not indicative of future results.\n'
        '    </p>\n'
        '    <p class=\"disclaimer-subtext\">\n'
        '        This fund invests primarily in equity and related instruments. \n'
        '        Returns are not guaranteed and principal may be at risk.\n'
        '    </p>\n'
        '</div>',
        style='Intense Quote'
    )
    disclaimer_impl.paragraph_format.left_indent = Inches(0.5)
    for run in disclaimer_impl.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS['pass']

    doc.add_paragraph()

    impl2 = doc.add_paragraph()
    impl2.add_run("Locations Required:\n").font.bold = True
    locations = [
        "• fund-detail.html (before \"Invest Now\" button)",
        "• mutual-funds.html (on fund cards - abbreviated version)",
        "• basket-details.html (before checkout)",
        "• All investment confirmation pages",
    ]
    for loc in locations:
        impl2.add_run(loc + "\n")
    impl2.add_run("Priority: ").font.bold = True
    priority2 = impl2.add_run("CRITICAL - Week 1, Day 1-2")
    priority2.font.color.rgb = COLORS['critical']
    priority2.font.bold = True

    doc.add_paragraph()

    # Suggested CSS
    doc.add_heading('Suggested CSS Styling', 3)

    css_code = doc.add_paragraph(
        '.sebi-disclaimer-box {\n'
        '    background: linear-gradient(135deg, #FEF3C7 0%, #FDE68A 100%);\n'
        '    border-left: 4px solid #F59E0B;\n'
        '    border-radius: 12px;\n'
        '    padding: 16px;\n'
        '    margin: 20px 0;\n'
        '}\n\n'
        '.disclaimer-header {\n'
        '    display: flex;\n'
        '    align-items: center;\n'
        '    gap: 8px;\n'
        '    font-size: 14px;\n'
        '    font-weight: 600;\n'
        '    color: #92400E;\n'
        '    margin-bottom: 8px;\n'
        '}\n\n'
        '.risk-level {\n'
        '    display: flex;\n'
        '    align-items: center;\n'
        '    gap: 8px;\n'
        '    margin-bottom: 12px;\n'
        '    font-size: 13px;\n'
        '}\n\n'
        '.risk-level.moderate strong {\n'
        '    color: #F59E0B;\n'
        '    font-weight: 700;\n'
        '}\n\n'
        '.disclaimer-text {\n'
        '    font-size: 13px;\n'
        '    line-height: 1.5;\n'
        '    color: #78350F;\n'
        '    margin: 8px 0;\n'
        '}\n\n'
        '.disclaimer-subtext {\n'
        '    font-size: 12px;\n'
        '    color: #92400E;\n'
        '    margin-top: 8px;\n'
        '}',
        style='Intense Quote'
    )
    css_code.paragraph_format.left_indent = Inches(0.5)
    for run in css_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    doc.add_paragraph()

    # Action Items
    doc.add_heading('Action Items for SEBI Compliance', 2)

    sebi_actions = [
        "Obtain SEBI registration number from legal/compliance team (Day 1, 1 hour)",
        "Add SEBI registration footer to all 39 pages (Week 1, 6 hours)",
        "Create SEBI disclaimer component with proper styling (Week 1, 3 hours)",
        "Add market risks disclaimer to fund-detail.html (Week 1, 2 hours)",
        "Add abbreviated disclaimer to all fund cards in mutual-funds.html (Week 1, 3 hours)",
        "Add disclaimer to basket-details.html before investment (Week 1, 2 hours)",
        "Create risk categorization badges (Low/Moderate/High) (Week 2, 4 hours)",
        "Add \"Past performance not indicative of future results\" to returns (Week 2, 2 hours)",
        "Add grievance redressal contact information to footer (Week 2, 2 hours)",
        "Legal review of all disclosures for compliance (Week 2, 4 hours)",
    ]

    for i, action in enumerate(sebi_actions, 1):
        p = doc.add_paragraph(f"☐ ACTION {i}: {action}", style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

    # 3.2 Apple Financial App Requirements
    doc.add_heading('3.2 Apple Financial App Guidelines (3.2.1)', 1)

    apple_fin = doc.add_paragraph()
    apple_fin.add_run("Apple Requirement: ").font.bold = True
    apple_fin.add_run(
        "Apps facilitating trading in financial instruments must be submitted by "
        "the financial institution providing the services OR must use a public API "
        "from the institution in compliance with their terms.\n"
    )
    apple_fin.add_run("Reference: ").font.bold = True
    apple_fin.add_run("https://developer.apple.com/app-store/review/guidelines/#financial-services\n")
    apple_fin.add_run("Rejection Risk: ").font.bold = True
    risk_apple = apple_fin.add_run("CRITICAL")
    risk_apple.font.color.rgb = COLORS['critical']
    risk_apple.font.bold = True

    doc.add_paragraph()

    doc.add_heading('Verification Checklist', 2)

    verification = [
        "☐ Is NiveshPe the licensed SEBI-registered financial institution?",
        "☐ If distributor/aggregator, which AMC/institution APIs are being used?",
        "☐ Are API usage agreements in place and compliant with institution T&C?",
        "☐ Documentation available to submit to Apple during review?",
        "☐ Proper business entity information in App Store Connect?",
        "☐ Financial licenses and registrations documented?",
    ]

    for item in verification:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    critical_note = doc.add_paragraph()
    critical_note.add_run("⚠️ CRITICAL: ").font.bold = True
    critical_note.add_run("Priority: ").font.bold = True
    note_priority = critical_note.add_run("CRITICAL - Must clarify before submission (Week 1)")
    note_priority.font.color.rgb = COLORS['critical']
    note_priority.font.bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        "Apple will request documentation proving either:\n"
        "1. Your company is the registered financial institution, OR\n"
        "2. Written authorization from the institution whose API you're using"
    )

    doc.add_paragraph()

    doc.add_heading('Required Documentation for App Store Review', 2)

    docs = [
        "SEBI registration certificate (ARN certificate)",
        "Business entity registration documents",
        "API usage agreements with AMCs/financial institutions",
        "Terms & Conditions of all integrated financial services",
        "Privacy policy specific to financial data handling",
        "Data security certifications (if available)",
    ]

    for doc_item in docs:
        p = doc.add_paragraph(f"• {doc_item}", style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_implementation_roadmap(doc):
    """Add implementation roadmap section"""
    doc.add_heading('SECTION 7: IMPLEMENTATION ROADMAP', 0)

    doc.add_paragraph(
        "This section provides a week-by-week breakdown of the recommended "
        "implementation timeline to achieve App Store compliance."
    )

    doc.add_paragraph()

    # Week 1-2: Critical
    doc.add_heading('WEEK 1-2: CRITICAL FIXES (Must Complete Before Submission)', 1)

    critical_box = doc.add_paragraph()
    critical_run = critical_box.add_run("Priority: CRITICAL")
    critical_run.font.bold = True
    critical_run.font.color.rgb = COLORS['critical']
    critical_box.add_run(" | Estimated Effort: 60-80 hours")

    doc.add_paragraph()

    week1_tasks = [
        ("Day 1", "Obtain SEBI registration number from legal team", "1 hour"),
        ("Day 1-2", "Add SEBI registration footer to all 39 pages", "6 hours"),
        ("Day 2-3", "Implement market risks disclaimer on all fund pages", "8 hours"),
        ("Day 3-5", "Add aria-label to all back buttons (36 pages)", "12 hours"),
        ("Day 3-5", "Update touch targets from 40px to 44pt minimum", "8 hours"),
        ("Day 5-7", "Add aria-labels to all Material Symbols icons", "8 hours"),
        ("Day 7-10", "Associate all form inputs with labels properly", "6 hours"),
        ("Day 8-10", "Add aria-labels to status bar elements", "2 hours"),
        ("Day 10-12", "Add grievance redressal information to footer", "2 hours"),
        ("Day 12-14", "Test all changes with VoiceOver on actual device", "8 hours"),
    ]

    week1_table = doc.add_table(rows=1, cols=3)
    week1_table.style = 'Light List Accent 1'

    w1_hdr = week1_table.rows[0].cells
    w1_hdr[0].text = 'Timeline'
    w1_hdr[1].text = 'Task'
    w1_hdr[2].text = 'Effort'

    for cell in w1_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for timeline, task, effort in week1_tasks:
        row = week1_table.add_row().cells
        row[0].text = timeline
        row[1].text = task
        row[2].text = effort

    doc.add_paragraph()

    milestone1 = doc.add_paragraph()
    milestone1.add_run("✓ MILESTONE: ").font.bold = True
    milestone1.add_run("Critical rejection risks eliminated. App can pass basic accessibility and compliance checks.")

    doc.add_page_break()

    # Week 3-4: High Priority
    doc.add_heading('WEEK 3-4: HIGH PRIORITY', 1)

    high_box = doc.add_paragraph()
    high_run = high_box.add_run("Priority: HIGH")
    high_run.font.bold = True
    high_run.font.color.rgb = COLORS['high']
    high_box.add_run(" | Estimated Effort: 50-70 hours")

    doc.add_paragraph()

    week3_tasks = [
        ("Week 3", "Audit md3-tokens.css and create px to rem conversion table", "4 hours"),
        ("Week 3", "Convert all typography tokens from px to rem units", "8 hours"),
        ("Week 3", "Test all pages at 100%, 150%, 200% text scaling", "12 hours"),
        ("Week 3", "Add privacy policy and data collection transparency", "8 hours"),
        ("Week 4", "Fix container heights/widths that break with larger text", "16 hours"),
        ("Week 4", "Implement max-width constraints to prevent overflow", "8 hours"),
        ("Week 4", "Add CSS clamp() for responsive text scaling", "6 hours"),
        ("Week 4", "Add risk categorization badges to all funds", "6 hours"),
        ("Week 4", "Test with users requiring large text accessibility", "4 hours"),
    ]

    week3_table = doc.add_table(rows=1, cols=3)
    week3_table.style = 'Light List Accent 1'

    w3_hdr = week3_table.rows[0].cells
    w3_hdr[0].text = 'Timeline'
    w3_hdr[1].text = 'Task'
    w3_hdr[2].text = 'Effort'

    for cell in w3_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for timeline, task, effort in week3_tasks:
        row = week3_table.add_row().cells
        row[0].text = timeline
        row[1].text = task
        row[2].text = effort

    doc.add_paragraph()

    milestone2 = doc.add_paragraph()
    milestone2.add_run("✓ MILESTONE: ").font.bold = True
    milestone2.add_run("App supports Dynamic Type and has complete financial disclosures. Ready for comprehensive testing.")

    doc.add_paragraph()

    # Week 5-6: Medium Priority
    doc.add_heading('WEEK 5-6: MEDIUM PRIORITY & OPTIMIZATION', 1)

    medium_box = doc.add_paragraph()
    medium_run = medium_box.add_run("Priority: MEDIUM")
    medium_run.font.bold = True
    medium_run.font.color.rgb = COLORS['medium']
    medium_box.add_run(" | Estimated Effort: 30-50 hours")

    doc.add_paragraph()

    week5_tasks = [
        "Review and optimize all navigation patterns for iOS conventions",
        "Add auto-fill support to all relevant form fields",
        "Audit and optimize keyboard types for all inputs",
        "Implement reduced motion support for animations",
        "Performance optimization (load time <3s target)",
        "Add proper error state announcements for screen readers",
        "Review and fix any remaining layout issues",
        "Comprehensive cross-device testing (various iOS versions)",
    ]

    for task in week5_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_paragraph()

    # Week 7: Testing
    doc.add_heading('WEEK 7: FINAL TESTING & VALIDATION', 1)

    test_box = doc.add_paragraph()
    test_run = test_box.add_run("Priority: CRITICAL")
    test_run.font.bold = True
    test_run.font.color.rgb = COLORS['critical']
    test_box.add_run(" | Estimated Effort: 20-30 hours")

    doc.add_paragraph()

    week7_tasks = [
        "Complete VoiceOver testing on physical devices (all pages)",
        "Dynamic Type testing at all size categories",
        "Touch target verification with accessibility tools",
        "Color blindness simulation testing",
        "Reduced motion testing",
        "Performance testing on older devices",
        "Legal review of all SEBI disclosures",
        "Final pre-submission checklist completion",
        "App Store Connect metadata preparation",
        "Screenshots with accessibility features enabled",
    ]

    for task in week7_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_paragraph()

    final_milestone = doc.add_paragraph()
    final_milestone.add_run("✓ FINAL MILESTONE: ").font.bold = True
    final_run = final_milestone.add_run("App is App Store submission ready!")
    final_run.font.color.rgb = COLORS['pass']
    final_run.font.bold = True

    doc.add_page_break()

def add_testing_checklists(doc):
    """Add comprehensive testing checklists"""
    doc.add_heading('SECTION 8: TESTING CHECKLISTS', 0)

    doc.add_paragraph(
        "Use these checklists to verify compliance before App Store submission."
    )

    doc.add_paragraph()

    # VoiceOver Testing
    doc.add_heading('VoiceOver Testing Checklist', 1)

    voiceover_tests = [
        "☐ Enable VoiceOver on iOS device (Settings > Accessibility > VoiceOver)",
        "☐ Test navigation through all 39 pages using only swipe gestures",
        "☐ Verify all buttons announce their purpose clearly",
        "☐ Verify all links describe their destination",
        "☐ Test form submission flow completely hands-free",
        "☐ Verify error messages are announced properly",
        "☐ Test investment flow from fund discovery to confirmation",
        "☐ Verify all charts and graphs have text alternatives",
        "☐ Test authentication flow (OTP input with VoiceOver)",
        "☐ Verify status bar elements are accessible or hidden",
        "☐ Test with screen curtain enabled (completely blind simulation)",
        "☐ Use Xcode Accessibility Inspector for automated checks",
        "☐ Document any elements still lacking proper labels",
        "☐ Test with actual blind user if possible",
    ]

    for test in voiceover_tests:
        p = doc.add_paragraph(test, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Dynamic Type Testing
    doc.add_heading('Dynamic Type Testing Checklist', 1)

    dynamic_tests = [
        "☐ Go to Settings > Accessibility > Display & Text Size > Larger Text",
        "☐ Enable \"Larger Accessibility Sizes\"",
        "☐ Test at each size category:",
        "  ☐ XS (Extra Small)",
        "  ☐ S (Small)",
        "  ☐ M (Medium - default)",
        "  ☐ L (Large)",
        "  ☐ XL (Extra Large)",
        "  ☐ XXL (Extra Extra Large)",
        "  ☐ XXXL (Accessibility 1)",
        "  ☐ A2 (Accessibility 2)",
        "  ☐ A3 (Accessibility 3 - largest)",
        "☐ Verify no text is cut off at any size",
        "☐ Verify no overlapping elements",
        "☐ Verify buttons remain tappable",
        "☐ Verify horizontal scrolling when necessary",
        "☐ Test critical flows (authentication, investment) at largest size",
        "☐ Document any pages requiring layout adjustments",
    ]

    for test in dynamic_tests:
        p = doc.add_paragraph(test, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Touch Target Testing
    doc.add_heading('Touch Target Verification Checklist', 1)

    touch_tests = [
        "☐ Enable Accessibility Inspector in Xcode",
        "☐ Audit all interactive elements for 44×44pt minimum",
        "☐ Test back buttons on all pages",
        "☐ Test modal close buttons",
        "☐ Test icon-only action buttons",
        "☐ Test period selector buttons",
        "☐ Test tab bar items",
        "☐ Test all form input fields",
        "☐ Test with motor disability simulator (shake/tremor)",
        "☐ Verify adequate spacing between adjacent tap targets",
        "☐ Test on smallest supported device (iPhone SE)",
        "☐ Document any elements below minimum size",
    ]

    for test in touch_tests:
        p = doc.add_paragraph(test, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Financial Compliance
    doc.add_heading('Financial Compliance Verification', 1)

    compliance_tests = [
        "☐ SEBI registration number displayed on all 39 pages",
        "☐ Market risks disclaimer present on all fund detail pages",
        "☐ Risk categorization visible on all fund cards",
        "☐ Past performance disclaimer on returns display",
        "☐ Expense ratio disclosed for all funds",
        "☐ Exit load information available",
        "☐ Grievance redressal contact information accessible",
        "☐ Privacy policy link functional and accessible",
        "☐ Terms of Service link functional and accessible",
        "☐ Legal team review completed",
        "☐ All disclaimers use exact regulatory language",
        "☐ Font sizes meet minimum readability (12px+)",
    ]

    for test in compliance_tests:
        p = doc.add_paragraph(test, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Pre-Submission Final Checklist
    doc.add_heading('Pre-Submission Final Checklist', 1)

    final_header = doc.add_paragraph()
    final_run = final_header.add_run("⚠️ DO NOT SUBMIT until ALL items are checked ☑")
    final_run.font.bold = True
    final_run.font.color.rgb = COLORS['critical']
    final_run.font.size = Pt(14)

    doc.add_paragraph()

    final_tests = [
        "☐ All VoiceOver tests passing",
        "☐ All Dynamic Type tests passing",
        "☐ All touch targets verified ≥44×44pt",
        "☐ All SEBI disclosures present and verified by legal",
        "☐ Financial institution documentation ready for Apple",
        "☐ Privacy policy updated and accessible",
        "☐ App Privacy Details completed in App Store Connect",
        "☐ No crashes or critical bugs",
        "☐ Performance <3s load time on 4G network",
        "☐ Built with Xcode 15+ and iOS 17 SDK minimum",
        "☐ Supports latest iOS version",
        "☐ Screenshots taken (including with accessibility enabled)",
        "☐ App Store description mentions accessibility features",
        "☐ Metadata keywords optimized",
        "☐ Age rating appropriate",
        "☐ Category correct (Finance)",
        "☐ Support URL functional",
        "☐ Marketing URL functional",
        "☐ In-app purchases configured (if applicable)",
        "☐ Test flight beta testing completed",
        "☐ All team members reviewed and approved",
    ]

    for test in final_tests:
        p = doc.add_paragraph(test, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_code_templates(doc):
    """Add code templates section"""
    doc.add_heading('SECTION 9: CODE TEMPLATES & EXAMPLES', 0)

    doc.add_paragraph(
        "Ready-to-use code templates for implementing required fixes."
    )

    doc.add_paragraph()

    # Accessible Button Template
    doc.add_heading('Accessible Button Template', 1)

    doc.add_paragraph("Use this template for all icon-only buttons:")

    button_template = doc.add_paragraph(
        '<!-- ✓ COMPLIANT: Accessible button with proper aria-label -->\n'
        '<button \n'
        '    class=\"back-btn\"\n'
        '    onclick=\"history.back()\"\n'
        '    aria-label=\"Go back to previous page\">\n'
        '    <span class=\"material-symbols-outlined\" aria-hidden=\"true\">arrow_back</span>\n'
        '</button>\n\n'
        '/* CSS for 44×44pt minimum touch target */\n'
        '.back-btn {\n'
        '    width: 44px;  /* Minimum required */\n'
        '    height: 44px; /* Minimum required */\n'
        '    border-radius: 50%;\n'
        '    background: #F8F9FA;\n'
        '    border: none;\n'
        '    display: flex;\n'
        '    align-items: center;\n'
        '    justify-content: center;\n'
        '    cursor: pointer;\n'
        '    transition: all 0.2s ease;\n'
        '}\n\n'
        '.back-btn:active {\n'
        '    background: #E5E7EB;\n'
        '    transform: scale(0.95);\n'
        '}',
        style='Intense Quote'
    )
    button_template.paragraph_format.left_indent = Inches(0.5)
    for run in button_template.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    doc.add_paragraph()

    # SEBI Disclaimer Component
    doc.add_heading('SEBI Disclaimer Component', 1)

    doc.add_paragraph("Standard disclaimer for all fund pages:")

    sebi_template = doc.add_paragraph(
        '<!-- ✓ COMPLIANT: SEBI market risks disclaimer -->\n'
        '<div class=\"sebi-disclaimer-box\" role=\"alert\" aria-live=\"polite\">\n'
        '    <div class=\"disclaimer-header\">\n'
        '        <span class=\"material-symbols-outlined\" aria-hidden=\"true\">info</span>\n'
        '        <strong>Investment Risk Disclosure</strong>\n'
        '    </div>\n'
        '    \n'
        '    <!-- Risk Level Badge -->\n'
        '    <div class=\"risk-level moderate\">\n'
        '        <span>Risk Level: </span>\n'
        '        <strong>Moderately High</strong>\n'
        '    </div>\n'
        '    \n'
        '    <!-- Required SEBI Disclaimer -->\n'
        '    <p class=\"disclaimer-text\">\n'
        '        <strong>Mutual Funds are subject to market risks.</strong> \n'
        '        Please read all scheme related documents carefully before investing. \n'
        '        Past performance is not indicative of future results.\n'
        '    </p>\n'
        '    \n'
        '    <!-- Additional Fund-Specific Warning -->\n'
        '    <p class=\"disclaimer-subtext\">\n'
        '        This fund invests primarily in equity and related instruments. \n'
        '        Returns are not guaranteed and principal may be at risk.\n'
        '    </p>\n'
        '</div>',
        style='Intense Quote'
    )
    sebi_template.paragraph_format.left_indent = Inches(0.5)
    for run in sebi_template.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    doc.add_paragraph()

    # Footer Template
    doc.add_heading('Footer with SEBI Registration', 1)

    footer_template = doc.add_paragraph(
        '<!-- ✓ COMPLIANT: Footer with SEBI registration (add to ALL pages) -->\n'
        '<footer class=\"app-footer\">\n'
        '    <div class=\"footer-content\">\n'
        '        <!-- SEBI Registration -->\n'
        '        <p class=\"sebi-registration\">\n'
        '            <strong>NiveshPe</strong> is registered with SEBI as a \n'
        '            Mutual Fund Distributor (ARN-XXXXXX). Valid till: DD/MM/YYYY\n'
        '        </p>\n'
        '        \n'
        '        <!-- Legal Links -->\n'
        '        <p class=\"regulatory-links\">\n'
        '            <a href=\"/privacy-policy\" aria-label=\"Read our privacy policy\">Privacy Policy</a> | \n'
        '            <a href=\"/terms\" aria-label=\"Read terms of service\">Terms of Service</a> | \n'
        '            <a href=\"/grievances\" aria-label=\"Submit a grievance\">Grievance Redressal</a>\n'
        '        </p>\n'
        '        \n'
        '        <!-- Contact -->\n'
        '        <p class=\"footer-contact\">\n'
        '            For queries: <a href=\"mailto:support@niveshpe.com\">support@niveshpe.com</a>\n'
        '        </p>\n'
        '    </div>\n'
        '</footer>\n\n'
        '/* Footer CSS */\n'
        '.app-footer {\n'
        '    background: #F8F9FA;\n'
        '    border-top: 1px solid #E5E7EB;\n'
        '    padding: 20px;\n'
        '    margin-top: 40px;\n'
        '    font-size: 12px;\n'
        '    text-align: center;\n'
        '    color: #6B7280;\n'
        '}\n\n'
        '.sebi-registration {\n'
        '    font-weight: 600;\n'
        '    color: #1F2937;\n'
        '    margin-bottom: 12px;\n'
        '}\n\n'
        '.regulatory-links a {\n'
        '    color: #2563EB;\n'
        '    text-decoration: none;\n'
        '    padding: 8px; /* Adequate touch target */\n'
        '}\n\n'
        '.regulatory-links a:hover {\n'
        '    text-decoration: underline;\n'
        '}',
        style='Intense Quote'
    )
    footer_template.paragraph_format.left_indent = Inches(0.5)
    for run in footer_template.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    doc.add_page_break()

def add_appendices(doc):
    """Add appendices with references"""
    doc.add_heading('APPENDICES', 0)

    # Appendix A: Apple Citations
    doc.add_heading('Appendix A: Apple HIG & Review Guidelines Citations', 1)

    doc.add_paragraph("Official Apple Documentation References:")

    apple_refs = [
        ("Human Interface Guidelines - Accessibility", "https://developer.apple.com/design/human-interface-guidelines/accessibility"),
        ("HIG - Typography & Dynamic Type", "https://developer.apple.com/design/human-interface-guidelines/foundations/typography"),
        ("HIG - Layout & Touch Targets", "https://developer.apple.com/design/human-interface-guidelines/layout"),
        ("App Store Review Guidelines (Full)", "https://developer.apple.com/app-store/review/guidelines/"),
        ("Review Guidelines - Section 3.2.1 Financial Services", "https://developer.apple.com/app-store/review/guidelines/#financial-services"),
        ("Review Guidelines - Section 4.0 Design", "https://developer.apple.com/app-store/review/guidelines/#design"),
        ("Review Guidelines - Section 5.1 Privacy", "https://developer.apple.com/app-store/review/guidelines/#privacy"),
        ("Accessibility - VoiceOver Guide", "https://developer.apple.com/accessibility/voiceover/"),
        ("Xcode Accessibility Inspector", "https://developer.apple.com/documentation/accessibility/accessibility-inspector"),
    ]

    for title, url in apple_refs:
        p = doc.add_paragraph()
        p.add_run(f"• {title}\n").font.bold = True
        link_run = p.add_run(f"  {url}")
        link_run.font.color.rgb = COLORS['blue']
        link_run.font.underline = True

    doc.add_paragraph()

    # Appendix B: WCAG
    doc.add_heading('Appendix B: WCAG 2.1 Compliance Mapping', 1)

    doc.add_paragraph("Relevant WCAG 2.1 Level AA Guidelines:")

    wcag_refs = [
        ("1.3.1 Info and Relationships", "Proper HTML semantics, form labels", "FAIL"),
        ("1.4.3 Contrast (Minimum)", "4.5:1 for normal text, 3:1 for large", "PASS"),
        ("1.4.4 Resize Text", "Text can be resized to 200% without loss", "FAIL"),
        ("1.4.10 Reflow", "No 2D scrolling at 320px width", "PASS"),
        ("2.4.4 Link Purpose", "Link purpose clear from text/context", "FAIL"),
        ("2.5.5 Target Size", "Touch targets at least 44×44 CSS pixels", "PARTIAL"),
        ("4.1.2 Name, Role, Value", "All UI components properly identified", "FAIL"),
    ]

    wcag_table = doc.add_table(rows=1, cols=3)
    wcag_table.style = 'Light Grid Accent 1'

    wcag_hdr = wcag_table.rows[0].cells
    wcag_hdr[0].text = 'WCAG Guideline'
    wcag_hdr[1].text = 'Requirement'
    wcag_hdr[2].text = 'Status'

    for cell in wcag_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for guideline, requirement, status in wcag_refs:
        row = wcag_table.add_row().cells
        row[0].text = guideline
        row[1].text = requirement
        row[2].text = status

    doc.add_paragraph()

    wcag_link = doc.add_paragraph()
    wcag_link.add_run("Full WCAG 2.1 Guidelines: ").font.bold = True
    wcag_url = wcag_link.add_run("https://www.w3.org/WAI/WCAG21/Understanding/")
    wcag_url.font.color.rgb = COLORS['blue']
    wcag_url.font.underline = True

    doc.add_paragraph()

    # Appendix C: SEBI
    doc.add_heading('Appendix C: SEBI Regulations Summary', 1)

    doc.add_paragraph("Key SEBI Mutual Funds Regulations for Mobile Apps:")

    sebi_regs = [
        "SEBI (Mutual Funds) Regulations, 1996 - Base regulatory framework",
        "Regulation 52 - Advertisement Code (applies to app content)",
        "Regulation 59 - Disclosure requirements in offer documents",
        "SEBI Circular 2024 - Enhanced disclosure requirements for direct/regular plans",
        "Risk-o-meter color coding - Required from December 5, 2024",
        "Monthly portfolio disclosure - Within 10 days of month end",
        "Standard risk disclaimer - Mandatory on all promotional material",
    ]

    for reg in sebi_regs:
        doc.add_paragraph(f"• {reg}", style='List Bullet')

    doc.add_paragraph()

    sebi_link = doc.add_paragraph()
    sebi_link.add_run("SEBI Official Regulations: ").font.bold = True
    sebi_url = sebi_link.add_run("https://www.sebi.gov.in/legal/regulations.html")
    sebi_url.font.color.rgb = COLORS['blue']
    sebi_url.font.underline = True

    doc.add_paragraph()

    # Appendix D: Research Sources
    doc.add_heading('Appendix D: Research Sources & References', 1)

    doc.add_paragraph("This document was compiled using the following sources:")

    sources = [
        ("iOS App Store Review Guidelines", "https://appfollow.io/blog/app-store-review-guidelines"),
        ("Apple Developer Forums - Financial Apps", "https://developer.apple.com/forums/thread/118805"),
        ("iOS Accessibility Guidelines 2025", "https://medium.com/@david-auerbach/ios-accessibility-guidelines-best-practices-for-2025-6ed0d256200e"),
        ("App Store Compliance Guide", "https://www.appeneure.com/blog/how-to-ensure-app-store-compliance-for-ios-apps/"),
        ("SEBI Mutual Fund Regulations", "https://www.bajajfinserv.in/investments/sebi-regulations-for-mutual-funds"),
        ("App Store 2024 Statistics", "https://www.macrumors.com/2025/05/30/app-store-2024-transparency-report/"),
    ]

    for title, url in sources:
        p = doc.add_paragraph()
        p.add_run(f"• {title}\n").font.bold = True
        link_run = p.add_run(f"  {url}")
        link_run.font.color.rgb = COLORS['blue']
        link_run.font.underline = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Document metadata
    meta = doc.add_paragraph()
    meta.add_run("Document Version: ").font.bold = True
    meta.add_run("1.0\n")
    meta.add_run("Generated: ").font.bold = True
    meta.add_run("December 2024\n")
    meta.add_run("Last Updated: ").font.bold = True
    meta.add_run("December 9, 2024\n")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

def save_document(doc, filename):
    """Save the document"""
    doc.save(filename)
    file_size = os.path.getsize(filename) / (1024 * 1024)  # MB
    return file_size

def main():
    """Main execution function"""
    print("Generating iOS App Store Compliance Document...")
    print("=" * 60)

    # Create document
    doc = create_document()
    print("✓ Document created")

    # Add all sections
    add_title_page(doc)
    print("✓ Title page added")

    add_executive_summary(doc)
    print("✓ Executive summary added")

    add_accessibility_section(doc)
    print("✓ Accessibility section added")

    add_financial_compliance_section(doc)
    print("✓ Financial compliance section added")

    add_implementation_roadmap(doc)
    print("✓ Implementation roadmap added")

    add_testing_checklists(doc)
    print("✓ Testing checklists added")

    add_code_templates(doc)
    print("✓ Code templates added")

    add_appendices(doc)
    print("✓ Appendices added")

    # Save document
    filename = "iOS_App_Store_Compliance_Audit_NiveshPe.docx"
    file_size = save_document(doc, filename)

    print("=" * 60)
    print(f"✓ Document saved: {filename}")
    print(f"✓ File size: {file_size:.2f} MB")
    print(f"✓ Location: {os.path.abspath(filename)}")
    print("=" * 60)
    print("\nDocument generation complete!")
    print("\nNext steps:")
    print("1. Open the document in Microsoft Word")
    print("2. Review all sections for accuracy")
    print("3. Share with development and legal teams")
    print("4. Begin Week 1 critical fixes immediately")

if __name__ == "__main__":
    main()
